# -*- coding: utf-8 -*-
"""向 Word 文档插入 Office Math（OMML）公式，在 Word 中显示效果接近公式编辑器 / MathType。"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from mathml2omml import convert as mathml_to_omml

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def add_display_equation(document: Document, mathml: str, caption: str = "") -> None:
    """居中插入一行 OMML 显示公式。"""
    omml_inner = mathml_to_omml(mathml.strip())
    if omml_inner.startswith("<m:oMath"):
        omath_block = omml_inner
    else:
        omath_block = f"<m:oMath>{omml_inner}</m:oMath>"

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omath_para_xml = (
        f'<m:oMathPara xmlns:m="{MATH_NS}">'
        f"{omath_block}"
        "</m:oMathPara>"
    )
    paragraph._p.append(parse_xml(omath_para_xml))

    if caption:
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_paragraph.add_run(caption)
        run.italic = True


# 常用公式 MathML（Word 打开后可右键“专业型”进一步美化）
EQUATION_ZSCORE = r"""
<math xmlns="http://www.w3.org/1998/Math/MathML">
  <mrow>
    <msub><mi>z</mi><mi>i</mi></msub>
    <mo>=</mo>
    <mfrac>
      <mrow><msub><mi>x</mi><mi>i</mi></msub><mo>−</mo><msub><mi>μ</mi><mi>queue</mi></msub></mrow>
      <msub><mi>σ</mi><mi>queue</mi></msub>
    </mfrac>
  </mrow>
</math>
"""

EQUATION_LOGISTIC_LEVEL = r"""
<math xmlns="http://www.w3.org/1998/Math/MathML">
  <mrow>
    <mi>E</mi><mo stretchy="false">[</mo>
    <msub><mi>L</mi><mi>k</mi></msub>
    <mo stretchy="false">|</mo>
    <mi>s</mi><mo stretchy="false">]</mo>
    <mo>=</mo>
    <msub><mi>b</mi><mi>k</mi></msub>
    <mo>+</mo>
    <mfrac>
      <mn>1</mn>
      <mrow>
        <mn>1</mn><mo>+</mo>
        <mi>exp</mi><mo>⁡</mo>
        <mo stretchy="false">(</mo>
        <mo>−</mo>
        <mfrac>
          <mrow><mi>s</mi><mo>−</mo><msub><mi>τ</mi><mi>k</mi></msub></mrow>
          <msub><mi>w</mi><mi>k</mi></msub>
        </mfrac>
        <mo stretchy="false">)</mo>
      </mrow>
    </mfrac>
    <mo stretchy="false">(</mo>
    <msub><mi>a</mi><mi>k</mi></msub>
    <mo>−</mo>
    <msub><mi>b</mi><mi>k</mi></msub>
    <mo stretchy="false">)</mo>
  </mrow>
</math>
"""

EQUATION_PRIOR_STAGE = r"""
<math xmlns="http://www.w3.org/1998/Math/MathML">
  <mrow>
    <mi>P</mi><mo stretchy="false">(</mo>
    <mi>H</mi>
    <mo stretchy="false">|</mo>
    <mtext>stage</mtext>
    <mo>=</mo>
    <mi>s</mi>
    <mo stretchy="false">)</mo>
    <mo>≈</mo>
    <msub><mi>p</mi><mi>s</mi></msub>
    <mo>·</mo>
    <msub><mi>λ</mi><mi>s</mi></msub>
  </mrow>
</math>
"""

EQUATION_ODDS = r"""
<math xmlns="http://www.w3.org/1998/Math/MathML">
  <mrow>
    <mtext>odds</mtext><mo stretchy="false">(</mo><mi>H</mi><mo stretchy="false">)</mo>
    <mo>=</mo>
    <mfrac>
      <mrow><mi>P</mi><mo stretchy="false">(</mo><mi>H</mi><mo stretchy="false">)</mo></mrow>
      <mrow><mn>1</mn><mo>−</mo><mi>P</mi><mo stretchy="false">(</mo><mi>H</mi><mo stretchy="false">)</mo></mrow>
    </mfrac>
  </mrow>
</math>
"""

EQUATION_ODDS_UPDATE = r"""
<math xmlns="http://www.w3.org/1998/Math/MathML">
  <mrow>
    <mtext>odds</mtext><mo stretchy="false">(</mo>
    <mi>H</mi><mo stretchy="false">|</mo><mi>D</mi>
    <mo stretchy="false">)</mo>
    <mo>=</mo>
    <mtext>odds</mtext><mo stretchy="false">(</mo><mi>H</mi><mo stretchy="false">)</mo>
    <mo>×</mo>
    <munderover>
      <mo>∏</mo>
      <mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow>
      <mi>n</mi>
    </munderover>
    <mrow>
      <mi>L</mi><mi>R</mi><mo stretchy="false">(</mo>
      <msub><mi>z</mi><mi>i</mi></msub>
      <mo stretchy="false">)</mo>
    </mrow>
  </mrow>
</math>
"""

EQUATION_POSTERIOR = r"""
<math xmlns="http://www.w3.org/1998/Math/MathML">
  <mrow>
    <mi>P</mi><mo stretchy="false">(</mo>
    <mi>H</mi><mo stretchy="false">|</mo><mi>D</mi>
    <mo stretchy="false">)</mo>
    <mo>=</mo>
    <mfrac>
      <mrow>
        <mtext>odds</mtext><mo stretchy="false">(</mo>
        <mi>H</mi><mo stretchy="false">|</mo><mi>D</mi>
        <mo stretchy="false">)</mo>
      </mrow>
      <mrow>
        <mn>1</mn><mo>+</mo>
        <mtext>odds</mtext><mo stretchy="false">(</mo>
        <mi>H</mi><mo stretchy="false">|</mo><mi>D</mi>
        <mo stretchy="false">)</mo>
      </mrow>
    </mfrac>
  </mrow>
</math>
"""

EQUATION_LR_Z_POS = r"""
<math xmlns="http://www.w3.org/1998/Math/MathML">
  <mrow>
    <mi>L</mi><mi>R</mi><mo stretchy="false">(</mo>
    <msub><mi>z</mi><mi>i</mi></msub>
    <mo stretchy="false">)</mo>
    <mo>=</mo>
    <msubsup>
      <mrow><mi>L</mi><mi>R</mi></mrow>
      <mn>0</mn>
      <mrow><mo>min</mo><mo stretchy="false">(</mo><mrow><mo>|</mo><msub><mi>z</mi><mi>i</mi></msub><mo>|</mo></mrow><mo>,</mo><msub><mi>z</mi><mtext>cap</mtext></msub><mo stretchy="false">)</mo></mrow>
    </msubsup>
    <mo>,</mo>
    <msub><mi>z</mi><mi>i</mi></msub>
    <mo>&gt;</mo>
    <mn>0</mn>
  </mrow>
</math>
"""

EQUATION_LR_Z_NEG = r"""
<math xmlns="http://www.w3.org/1998/Math/MathML">
  <mrow>
    <mi>L</mi><mi>R</mi><mo stretchy="false">(</mo>
    <msub><mi>z</mi><mi>i</mi></msub>
    <mo stretchy="false">)</mo>
    <mo>=</mo>
    <msubsup>
      <mrow><mi>L</mi><mi>R</mi></mrow>
      <mn>0</mn>
      <mrow><mo>−</mo><mo>min</mo><mo stretchy="false">(</mo><mrow><mo>|</mo><msub><mi>z</mi><mi>i</mi></msub><mo>|</mo></mrow><mo>,</mo><msub><mi>z</mi><mtext>cap</mtext></msub><mo stretchy="false">)</mo></mrow>
    </msubsup>
    <mo>,</mo>
    <msub><mi>z</mi><mi>i</mi></msub>
    <mo>&lt;</mo>
    <mn>0</mn>
  </mrow>
</math>
"""

EQUATION_LOG_LIKELIHOOD = r"""
<math xmlns="http://www.w3.org/1998/Math/MathML">
  <mrow>
    <mi>log</mi><mo>⁡</mo>
    <mi>P</mi><mo stretchy="false">(</mo>
    <mi>D</mi><mo stretchy="false">|</mo><mi>s</mi><mo stretchy="false">)</mo>
    <mo>∝</mo>
    <mo>−</mo>
    <mfrac><mn>1</mn><mn>2</mn></mfrac>
    <munder>
      <mo>∑</mo>
      <mi>k</mi>
    </munder>
    <msup>
      <mrow>
        <mo stretchy="false">(</mo>
        <mfrac>
          <mrow>
            <msub><mi>x</mi><mi>k</mi></msub>
            <mo>−</mo>
            <msub><mi>E</mi><mi>k</mi></msub>
          </mrow>
          <msub><mi>σ</mi><mi>k</mi></msub>
        </mfrac>
        <mo stretchy="false">)</mo>
      </mrow>
      <mn>2</mn>
    </msup>
  </mrow>
</math>
"""
