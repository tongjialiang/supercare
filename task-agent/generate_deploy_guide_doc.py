#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成简洁版《系统部署与运行说明》Word 文档。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

PROJECT_ROOT = Path("/srv/supercare")
DEPLOY_DOC = PROJECT_ROOT / "比赛文档/3_系统部署与运行说明_超级[GP-护士-照护员]智能体协同算法.docx"
DELIVERABLE_DOC = PROJECT_ROOT / "比赛文档/交付物" / DEPLOY_DOC.name


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for col_index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[col_index], header, bold=True)
    for row_index, row_values in enumerate(rows, start=1):
        for col_index, value in enumerate(row_values):
            _set_cell_text(table.rows[row_index].cells[col_index], value)
    doc.add_paragraph()


def _add_code_block(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    paragraph.paragraph_format.space_after = Pt(6)


def build_deploy_guide() -> Document:
    doc = Document()

    title = doc.add_heading("系统部署与运行说明", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("超级 [GP - 护士 - 照护员] 智能体协同算法")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading("一、代码开源地址", level=1)
    doc.add_paragraph("项目代码托管于 GitHub：")
    _add_table(
        doc,
        ["字段", "值"],
        [
            ["GitHub 仓库", "<请填写实际开源仓库地址>"],
            ["默认分支", "main"],
            ["项目根目录", "/srv/supercare"],
        ],
    )

    doc.add_heading("二、环境准备", level=1)
    doc.add_paragraph(
        "系统为 Python 3.10+ 项目，四大模块按流水线顺序运行："
        "DataAgent1 → DataAgent2 → 序贯分期与贝叶斯 → 超级 GP 协同。"
    )
    doc.add_paragraph("一次性安装依赖：")
    _add_code_block(
        doc,
        "cd /srv/supercare\n"
        "python3 -m pip install -r requirements-data-agent1.txt -r requirements-quickstart.txt",
    )
    doc.add_paragraph(
        "必填配置：编辑 config/data_agent1_config.json，填入 qwen.api_key（通义千问 DashScope）。"
    )
    doc.add_paragraph(
        "示例数据已内置：DataSource/陈女士/ms_chen_CareCase.xlsx，"
        "以及 DataSource/ 下 10 例机构队列病例。"
    )

    doc.add_heading("三、快速运行", level=1)

    doc.add_heading("3.1 推荐：FastAPI 一键运行", level=2)
    doc.add_paragraph("启动服务：")
    _add_code_block(
        doc,
        "cd /srv/supercare\n"
        "uvicorn supercare_quickstart_api:app --host 0.0.0.0 --port 8765",
    )
    doc.add_paragraph(
        "浏览器打开 http://127.0.0.1:8765/docs，调用 POST /run/pipeline/all。"
        "推荐参数：fast_mode=true（约 1–2 分钟解析）、demo_task_agent=true（仅演示 A3 超级 GP）。"
        "返回 job_id 后，用 GET /jobs/{job_id} 轮询直至 status 为 success。"
    )
    doc.add_paragraph("常用辅助接口：GET /health、GET /config/check、GET /outputs。")

    doc.add_heading("3.2 备选：命令行四步", level=2)
    _add_code_block(
        doc,
        "# ① 长者健康计算图 DataAgent\n"
        "python /srv/supercare/data_agent_task1.py \\\n"
        '  --input "/srv/supercare/DataSource/陈女士/ms_chen_CareCase.xlsx" \\\n'
        '  --output-root "/srv/supercare/data-agent1" \\\n'
        '  --config "/srv/supercare/config/data_agent1_config.json" \\\n'
        "  --parser-mode openpyxl\n\n"
        "# ② 三超循证知识基座 DataAgent\n"
        "python /srv/supercare/data_agent_task2.py \\\n"
        '  --input-graph "/srv/supercare/data-agent1/ms_chen_CareCase/results/ms_chen_CareCase_elder_health_computing_graph.json" \\\n'
        '  --output-root "/srv/supercare/data-agent2" \\\n'
        '  --config "/srv/supercare/config/data_agent1_config.json"\n\n'
        "# ③ 序贯分期与贝叶斯融合推断\n"
        "python /srv/supercare/task-agent/run_staging_pipeline.py\n"
        "python /srv/supercare/task-agent/generate_a0_visualizations.py\n\n"
        "# ④ 超级 GP 协同（演示）\n"
        "python /srv/supercare/task-agent/test_agent_a3_cda_hub.py",
    )
    doc.add_paragraph(
        "全量 A1–A14 协同：python /srv/supercare/task-agent/run_agents_a1_a14.py（耗时较长）。"
    )

    doc.add_heading("四、目录说明", level=1)
    _add_table(
        doc,
        ["目录 / 文件", "功能", "关键产物"],
        [
            ["DataSource/", "示例病例 Excel 与机构队列原始数据", "ms_chen_CareCase.xlsx 等"],
            ["config/", "通义千问 API、MinerU 等运行配置", "data_agent1_config.json"],
            ["data_agent_task1.py", "模块一：长者健康计算图 DataAgent 入口", "elder_health_computing_graph.json"],
            ["data-agent1/", "DataAgent1 运行输出目录", "data-agent1/<病例>/results/*.json"],
            ["data_agent_task2.py", "模块二：三超循证知识基座 DataAgent 入口", "六份 JSONL 语料"],
            ["data-agent2/", "DataAgent2 运行输出与质量对比", "batch_summary.json"],
            ["task-agent/", "序贯分期、贝叶斯推断、A1–A15 协同智能体", "output/a0_*.json、output/a3_*.pdf"],
            ["task-agent/staging/", "A0 序贯分期与贝叶斯核心算法", "pipeline.py 等"],
            ["supercare_quickstart_api.py", "评委一键运行 FastAPI 服务", "POST /run/pipeline/all"],
            ["比赛文档/", "技术报告、部署说明、测试日志等交付文档", "1_技术报告、3_部署说明、4_运行日志"],
            ["比赛文档/交付物/", "答辩与评审用精简交付包", "快速运行指南.md 等"],
        ],
    )

    doc.add_heading("五、运行验证", level=1)
    _add_table(
        doc,
        ["模块", "关键产物"],
        [
            ["DataAgent1", "data-agent1/.../ms_chen_CareCase_elder_health_computing_graph.json"],
            ["DataAgent2", "data-agent2/batch_summary.json"],
            ["序贯—贝叶斯", "task-agent/output/a0_序贯疾病分期结果.json、a0_贝叶斯风险后验.json"],
            ["协同算法", "task-agent/output/a3_GP协作专业答复.pdf"],
        ],
    )
    doc.add_paragraph("FastAPI 模式下可执行：curl -s http://127.0.0.1:8765/outputs | python3 -m json.tool")

    doc.add_heading("六、相关文档", level=1)
    doc.add_paragraph("详细测试与对比实验见：比赛文档/4_系统运行日志与对比实验_....docx")
    doc.add_paragraph("逐步操作见：比赛文档/快速运行指南.md")

    return doc


def main() -> None:
    doc = build_deploy_guide()
    DEPLOY_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DEPLOY_DOC))
    DELIVERABLE_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DELIVERABLE_DOC))
    print(f"已生成: {DEPLOY_DOC}")
    print(f"已同步: {DELIVERABLE_DOC}")


if __name__ == "__main__":
    main()
