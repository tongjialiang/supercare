#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""由《快速运行指南.md》结构生成 Word 版本。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PROJECT_ROOT = Path("/srv/supercare")
GUIDE_DOC = PROJECT_ROOT / "比赛文档/快速运行指南.docx"
DELIVERABLE_DOC = PROJECT_ROOT / "比赛文档/交付物/快速运行指南.docx"
ASSETS_DIR = PROJECT_ROOT / "比赛文档/素材"
API_BASE = "http://corp.tunnel.izhaohu.com:8765"
API_DOCS = f"{API_BASE}/docs"
LOCAL_DOCS = "http://127.0.0.1:8765/docs"
APPLICANT_UNIT = "上海爱照护医疗科技有限公司"

FIGURE_ITEMS = [
    ("图1_BPSD返院全周期智能体工作流.png", "图 1  BPSD 返院/返机构全周期闭环智能体工作流（协同层架构）"),
    ("图2_照护科学实践图谱.png", "图 2  照护科学实践图谱（三层核心智能体协同的干预与成效评价机制）"),
    ("图3_疾病进展曲线_陈女士.png", "图 3  疾病进展曲线与分期对比（推断层：序贯分期 + MCMC，示例病例：陈女士）"),
    ("图4_A0序贯分期与贝叶斯推断架构.png", "图 4  A0 本地计算流水线（数据层 → 推断层 → 工具层 → 协同智能体）"),
]


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for col_index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[col_index], header, bold=True)
    for row_index, row_values in enumerate(rows, start=1):
        for col_index, value in enumerate(row_values):
            _set_cell_text(table.rows[row_index].cells[col_index], value)
    doc.add_paragraph()


def _add_code_block(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    paragraph.paragraph_format.space_after = Pt(6)


def _add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _add_faq(doc: Document, question: str, answer: str) -> None:
    paragraph = doc.add_paragraph()
    question_run = paragraph.add_run(f"Q：{question}")
    question_run.bold = True
    doc.add_paragraph(f"A：{answer}")


def _add_body_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)


def _add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.is_file():
        _add_body_paragraph(doc, f"[图片缺失] {caption}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _add_research_overview(doc: Document) -> None:
    """文档开头：申报单位、研究内容与配图。"""
    unit = doc.add_paragraph()
    unit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    unit_run = unit.add_run(f"项目申报单位：{APPLICANT_UNIT}")
    unit_run.bold = True
    unit_run.font.name = "宋体"
    unit_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.add_paragraph()

    doc.add_heading("研究内容概述", level=1)
    _add_body_paragraph(
        doc,
        "提出一套面向养老照护场景、以「数据—知识—推断—协同」四层贯通架构为核心的全链路智能体解决方案。"
        "以长期照护「证据+判断」双轮驱动为总体方法，构建覆盖「基态—触发—就医—返回」全周期的智能照护闭环。",
    )

    doc.add_heading("数据层", level=2)
    _add_body_paragraph(
        doc,
        "提出基于 MinerU 的长者健康智能计算图 DataAgent 技术：依托多智能体协同工作流与工具链，"
        "构建「洞察—证据」双层图结构，将多源异构的健康数据转化为可计算、可推理、可追溯的长者健康计算图。"
        "消融实验显示，MinerU 方案在图节点、图边与语义完整性方面优于常规解析方式。",
    )

    doc.add_heading("知识层", level=2)
    _add_body_paragraph(
        doc,
        "提出三超循证知识基座 DataAgent：涵盖原始记录层、叙事重建层、过程—成效关联分析层、循证规则层等递进式语料加工体系，"
        "将零散照护记录沉淀为高质量循证知识。对比实验显示，相较于通用大模型基线，本方案结构化覆盖率提升 0.70、"
        "证据引用率提升 0.20、因果表达率提升 0.18，知识质量大幅优化。",
    )

    doc.add_heading("推断层", level=2)
    _add_body_paragraph(
        doc,
        "创新性提出序贯分期与贝叶斯风险融合推断算法：针对老年 BPSD，依托爱照护长期纵向照护数据，"
        "选取 NPI、MMSE/MoCA、Barthel 等核心评估指标，通过 logistic 曲线刻画病情演变规律，结合 MCMC 算法判定疾病分期；"
        "风险推断过程中，分层先验概率引自国内外权威文献，再结合长者个体指标标准化 Z 分数与队列均值差异完成贝叶斯似然比递推更新，"
        "输出长者 30 天内病情恶化的后验风险概率。同时设置返院 D2、D7、D30 多节点分级风险阈值，"
        "为照护升级、临床会诊提供可解释、可校准的量化决策依据。",
    )

    doc.add_heading("协同层", level=2)
    _add_body_paragraph(
        doc,
        "创新性提出超级「GP-护士-照护员」智能体协同算法：通过搭建总控智能体 + 3 类角色中枢 + 11 个任务智能体 + "
        "13 类标准化工具的协同架构，依托算法推断、多智能体协同与人在回路机制，"
        "实现 BPSD 患者照护过程中的事件识别、任务分配、阶段评估与成效复盘。"
        "同时持续沉淀标准化照护知识图谱与行业高价值语料，为养老照护经验传承、知识标准化与规模化落地提供核心技术支撑。"
        "A/B 测试证明，该工具链可使智能体输出完整度提升 23.04%，显著提升照护方案的结构化与可执行性。",
    )
    _add_body_paragraph(
        doc,
        "依托超级「GP-护士-照护员」智能体协同算法，完成五类养老照护典型场景应用演示，"
        "包括 BPSD 患者住院后返回照护机构全流程、制定照护计划到执行全闭环、照护任务执行与质控、"
        "Dashboard 可视化管理、照护成效管理等，有效验证了所提方案的通用性与落地价值，"
        "能够显著提升养老照护的精细化与智能化水平。",
    )

    modules = doc.add_paragraph(
        "四大模块：长者智能计算图 DataAgent → 三超循证知识基座 DataAgent → "
        "序贯分期与贝叶斯风险融合推断算法 → 超级「GP-护士-照护员」智能体协同算法。"
    )
    modules.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("配图素材", level=2)
    for filename, caption in FIGURE_ITEMS:
        _add_figure(doc, ASSETS_DIR / filename, caption)


def _add_output_catalog_tables(doc: Document) -> None:
    """§4 产物说明（按模块）。"""
    doc.add_heading("长者智能计算图 DataAgent（2 个）", level=3)
    _add_table(
        doc,
        ["文件", "内容说明", "下载 key"],
        [
            [
                "健康计算图 JSON",
                "陈女士病例的长者健康计算图结构化数据，含节点、边、洞察与证据链，供下游模块读取。",
                "data_agent1_graph_json",
            ],
            [
                "健康计算图 PNG",
                "健康计算图可视化展示图，用于答辩演示节点类别、关联关系与关键洞察分布。",
                "data_agent1_graph_png",
            ],
        ],
    )

    doc.add_heading("三超循证知识基座 DataAgent（7 个）", level=3)
    _add_table(
        doc,
        ["文件", "内容说明", "下载 key"],
        [
            [
                "batch_summary.json",
                "六份训练语料的批量生成汇总，含各角色 CoT/非 CoT 语料的条数、质量指标与对比结论。",
                "data_agent2_batch_summary",
            ],
            [
                "super_doctor_with_cot_50.jsonl",
                "超级医生带思维链语料 50 条，含 L1–L4 四层循证轨迹与证据引用，供 GP 决策微调。",
                "da2_corpus_super_doctor_with_cot_50",
            ],
            [
                "super_doctor_without_cot_50.jsonl",
                "超级医生无思维链语料 50 条，仅保留问答对，用于对比 CoT 增益的消融实验。",
                "da2_corpus_super_doctor_without_cot_50",
            ],
            [
                "super_nurse_with_cot_50.jsonl",
                "超级护士带思维链语料 50 条，覆盖评估结论、护理调整与风险预警等场景。",
                "da2_corpus_super_nurse_with_cot_50",
            ],
            [
                "super_nurse_without_cot_50.jsonl",
                "超级护士无思维链语料 50 条，结构与医生语料对称，便于跨角色质量对比。",
                "da2_corpus_super_nurse_without_cot_50",
            ],
            [
                "super_caregiver_with_cot_50.jsonl",
                "超级照护员带思维链语料 50 条，聚焦日常照护动作、非药物干预与家属沟通。",
                "da2_corpus_super_caregiver_with_cot_50",
            ],
            [
                "super_caregiver_without_cot_50.jsonl",
                "超级照护员无思维链语料 50 条，用于验证四层注入对实操建议质量的影响。",
                "da2_corpus_super_caregiver_without_cot_50",
            ],
        ],
    )

    doc.add_heading("序贯分期与贝叶斯风险融合推断算法（4 个）", level=3)
    _add_table(
        doc,
        ["文件", "内容说明", "下载 key"],
        [
            [
                "a0_序贯事件进展图_陈女士.png",
                "SuStaIn/MCMC 序贯事件槽位与共同分期进展示意图，展示各生物标志物事件在时间轴上的排列。",
                "a0_staging_event_png",
            ],
            [
                "a0_生物标志物分期进展热图_陈女士.png",
                "多生物标志物纵向序贯进展热图，直观对比各指标在不同分期的异常程度变化。",
                "a0_biomarker_heatmap_png",
            ],
            [
                "a0_疾病进展曲线_陈女士.png",
                "以共同分期为横轴、归一化进展指数为纵轴的疾病进展曲线，呈现多标志物协同演变趋势。",
                "a0_disease_curve_png",
            ],
            [
                "a0_贝叶斯先验后验图_陈女士.png",
                "文献校准贝叶斯先验与观测似然更新后的后验对比图，展示各分期风险概率变化。",
                "a0_bayesian_prior_posterior_png",
            ],
        ],
    )

    doc.add_heading("超级「GP-护士-照护员」智能体协同算法（19 个）", level=3)
    doc.add_paragraph("A1–A15 协同智能体 PDF（17 个）", style="Intense Quote")
    _add_table(
        doc,
        ["文件", "内容说明", "下载 key"],
        [
            [
                "a1_照护员协作专业答复.pdf",
                "超级照护员（CAA Hub）输出的非药物干预与日常执行方案，含跳舞兴趣引导、防跌倒要点。",
                "taskagent_a1_照护员协作专业答复",
            ],
            [
                "a2_护士协作专业答复.pdf",
                "超级护士（CCA Hub）输出的临床评估与护理调整建议，衔接医生决策与一线执行。",
                "taskagent_a2_护士协作专业答复",
            ],
            [
                "a3_GP协作专业答复.pdf",
                "超级 GP（CDA Hub）综合确认意见，统筹多角色协同照护路径与风险管控。",
                "taskagent_a3_GP协作专业答复",
            ],
            [
                "a4_出院摘要结构化包.pdf",
                "Discharge Parser 将出院病历解析为结构化字段包，供后续智能体统一读取。",
                "taskagent_a4_出院摘要结构化包",
            ],
            [
                "a5_ECR状态对比报告.pdf",
                "入院、出院、返院三时点的 ECR（情绪-认知-行为）状态横向对比分析。",
                "taskagent_a5_ECR状态对比报告",
            ],
            [
                "a6_返院适应期专项任务包.pdf",
                "返院适应期任务拆解清单，明确照护员、护士在各日的可执行动作。",
                "taskagent_a6_返院适应期专项任务包",
            ],
            [
                "a7_BPSD事件报告.pdf",
                "BPSD 行为事件的时序记录与情境描述，为升级决策提供结构化输入。",
                "taskagent_a7_BPSD事件报告",
            ],
            [
                "a8_BPSD事件处理建议.pdf",
                "Escalation Agent 针对 BPSD 事件的即时处置方案与是否升级会诊的建议。",
                "taskagent_a8_BPSD事件处理建议",
            ],
            [
                "a9_照护员周重点任务.pdf",
                "照护员本周重点任务与操作步骤，源自护士评估结论的下钻执行版。",
                "taskagent_a9_照护员周重点任务",
            ],
            [
                "a9_护士评估结论与调整建议.pdf",
                "护士周评估结论、生命体征与 ADL 变化摘要及护理计划调整建议。",
                "taskagent_a9_护士评估结论与调整建议",
            ],
            [
                "a10_三时点功能对比图.pdf",
                "住院前、出院、返院三个时间节点的 ADL/NPI 等功能指标可视化对比。",
                "taskagent_a10_三时点功能对比图",
            ],
            [
                "a11_返院一周综合报告.pdf",
                "返院第 7 天综合状态复盘，汇总行为、功能、体征与家属反馈。",
                "taskagent_a11_返院一周综合报告",
            ],
            [
                "a12_精神科远程会诊记录.pdf",
                "精神科远程会诊意见，含 BPSD 前驱信号识别与非药物干预优化建议。",
                "taskagent_a12_精神科远程会诊记录",
            ],
            [
                "a13_D30再稳定评估报告.pdf",
                "返院第 30 天再稳定期评估，判断 BPSD 是否进入低复发风险稳定期。",
                "taskagent_a13_D30再稳定评估报告",
            ],
            [
                "a13_照护成效中期报告.pdf",
                "返院 30 天照护成效中期总结，量化 NPI/ADL 改善与不良事件控制情况。",
                "taskagent_a13_照护成效中期报告",
            ],
            [
                "a14_TaskAgent审计评估报告.pdf",
                "A1–A14 全链路执行审计报告，记录各智能体调用顺序、产出与质量评分。",
                "taskagent_a14_TaskAgent审计评估报告",
            ],
            [
                "a15_老年健康序贯分期与贝叶斯风险报告.pdf",
                "A0 序贯分期与贝叶斯风险融合的综合算法报告，含文献先验与后验解读。",
                "taskagent_a15_老年健康序贯分期与贝叶斯风险报告",
            ],
        ],
    )
    doc.add_paragraph("语料与图谱（2 个）", style="Intense Quote")
    _add_table(
        doc,
        ["文件", "内容说明", "下载 key"],
        [
            [
                "a13_D30高价值语料.jsonl",
                "Outcome Report Agent 从 D30 照护闭环中提炼的高价值问答语料，用于后续模型微调与成效复盘。",
                "a13_d30_corpus_jsonl",
            ],
            [
                "照护科学实践图谱.jpg",
                "从 A1–A14 实践 PDF 凝练的照护科学实践图谱可视化，以超级照护员、护士、GP 为三认知核心向外辐射。",
                "care_science_practice_graph_jpg",
            ],
        ],
    )


def build_quickstart_guide() -> Document:
    doc = Document()

    title = doc.add_heading("SuperCare 快速运行指南", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_research_overview(doc)

    # §1 快速开始
    doc.add_heading("1. 快速开始（FastAPI）", level=1)
    doc.add_paragraph(
        f"服务启动后，打开 {API_DOCS}（本机亦可：{LOCAL_DOCS}），在 Swagger 页面直接调用接口即可。"
    )

    doc.add_heading("1.1 三步跑通全链路", level=2)
    _add_table(
        doc,
        ["步骤", "接口", "说明"],
        [
            ["①", "POST /run/pipeline/all", "一键启动四大模块，默认参数直接 Execute"],
            ["②", "GET /jobs/{job_id}", "用返回的 job_id 轮询，直到 status 为 success"],
            ["③", "GET /outputs/links", "获取全部产物下载链接"],
        ],
    )
    doc.add_paragraph("POST /run/pipeline/all 默认参数（答辩演示推荐）：")
    _add_table(
        doc,
        ["参数", "默认值", "说明"],
        [
            ["fast_mode", "true", "DataAgent1 用 openpyxl 加速（约 1–2 分钟）"],
            ["skip_task_agent", "false", "跑完整四大模块"],
            ["demo_task_agent", "true", "协同层仅演示 A3 超级 GP（约 1–3 分钟）"],
        ],
    )

    doc.add_heading("1.2 接口一览", level=2)
    doc.add_paragraph("概览", style="Intense Quote")
    _add_table(doc, ["接口", "作用"], [["GET /", "服务概览：四大模块名称、推荐入口路径"]])

    doc.add_paragraph("配置与状态", style="Intense Quote")
    _add_table(
        doc,
        ["接口", "作用"],
        [
            ["GET /health", "健康检查，确认服务在线"],
            ["GET /config/check", "检查配置文件与示例数据是否就绪（可选）"],
        ],
    )

    doc.add_paragraph("产物查看", style="Intense Quote")
    _add_table(
        doc,
        ["接口", "作用"],
        [
            ["GET /outputs", "快速检查 6 个核心验收文件是否已生成"],
            ["GET /outputs/links", "获取全部产物（约 32 个）的下载 URL 与中文说明"],
            ["GET /outputs/download?key=...", "按 key 下载单个产物文件"],
        ],
    )

    doc.add_paragraph("任务查询", style="Intense Quote")
    _add_table(
        doc,
        ["接口", "作用"],
        [
            ["GET /jobs", "列出最近后台任务（默认 20 条）"],
            ["GET /jobs/{job_id}", "查询单个任务进度、子步骤状态与日志摘要"],
        ],
    )

    doc.add_paragraph("模块运行（分步调试）", style="Intense Quote")
    _add_table(
        doc,
        ["接口", "作用"],
        [
            ["POST /run/data-agent1", "单独运行长者智能计算图 DataAgent"],
            ["POST /run/data-agent2", "单独运行三超循证知识基座 DataAgent"],
            ["POST /run/staging", "序贯分期与贝叶斯风险融合推断（纯本地）"],
            ["POST /run/staging/visualizations", "生成 A0 序贯分期配图（4 张 PNG）"],
            ["POST /run/task-agent/demo", "超级 GP（A3）演示"],
            ["POST /run/task-agent/full", "A1–A14 全量协同（耗时较长）"],
        ],
    )

    doc.add_paragraph("全链路运行", style="Intense Quote")
    _add_table(
        doc,
        ["接口", "作用"],
        [["POST /run/pipeline/all", "推荐入口：后台顺序执行四大模块"]],
    )

    doc.add_heading("1.3 下载产物示例", level=2)
    _add_code_block(
        doc,
        f"# 查看全部下载链接\n"
        f"curl -s {API_BASE}/outputs/links | python3 -m json.tool\n\n"
        f"# 下载照护科学实践图谱\n"
        f'curl -OJ "{API_BASE}/outputs/download?key=care_science_practice_graph_jpg"',
    )
    doc.add_paragraph(
        "返回字段：files 为完整列表；by_module 按四大模块分组；每条含 key、label、description、download_url。"
        "未生成的文件 exists 为 false。"
    )

    doc.add_heading("1.4 预计耗时", level=2)
    _add_table(
        doc,
        ["模式", "耗时"],
        [
            ["全链路（默认参数）", "约 10–20 分钟"],
            ["仅 A0 序贯贝叶斯", "< 1 分钟"],
            ["A1–A14 全量协同", "视 LLM 响应，显著更长"],
        ],
    )

    # §2 本地部署
    doc.add_heading("2. 本地部署", level=1)
    doc.add_heading("2.1 安装依赖", level=2)
    _add_code_block(
        doc,
        "cd /srv/supercare\n"
        "python3 -m pip install -r requirements-data-agent1.txt -r requirements-quickstart.txt",
    )

    doc.add_heading("2.2 启动 API 服务（常驻运行，推荐）", level=2)
    doc.add_paragraph(
        "关终端、注销后服务仍保持运行；机器重启后会自动拉起（需执行一次安装）。"
        "勿使用前台 uvicorn（关终端即退出）。"
    )
    doc.add_paragraph("方式 A：systemd 用户服务（推荐，开机自启 + 崩溃自动重启）", style="Intense Quote")
    _add_code_block(
        doc,
        "cd /srv/supercare\n"
        "bash scripts/install_quickstart_api_service.sh",
    )
    doc.add_paragraph("常用命令：")
    _add_code_block(
        doc,
        "systemctl --user status supercare-quickstart-api\n"
        "systemctl --user stop supercare-quickstart-api\n"
        "systemctl --user restart supercare-quickstart-api\n"
        "tail -f /srv/supercare/logs/quickstart_api.log",
    )
    doc.add_paragraph("方式 B：nohup 后台启动（简单，重启后需手动再执行）", style="Intense Quote")
    _add_code_block(
        doc,
        "cd /srv/supercare\n"
        "bash scripts/start_quickstart_api.sh    # 启动\n"
        "bash scripts/stop_quickstart_api.sh     # 停止",
    )
    doc.add_paragraph(f"浏览器打开 {API_DOCS}（本机：{LOCAL_DOCS}）即可调用。")

    doc.add_heading("2.3 配置说明（可选）", level=2)
    doc.add_paragraph("编辑 config/data_agent1_config.json：")
    _add_table(
        doc,
        ["配置项", "说明"],
        [
            ["qwen.api_key", "通义千问 DashScope API Key（DataAgent 与协同智能体 LLM 调用需要）"],
            ["mineru.api_url", "MinerU 远程解析地址（仅 fast_mode=false 正式模式需要）"],
        ],
    )
    doc.add_paragraph("默认 fast_mode=true 走 openpyxl 加速，可不配置 MinerU。")
    doc.add_paragraph("示例数据（已内置）：")
    _add_bullet(doc, "Excel：DataSource/陈女士/ms_chen_CareCase.xlsx")
    _add_bullet(doc, "数据队列：DataSource/ 下 10 例机构病例")

    doc.add_heading("2.4 命令行逐步运行（无 API）", level=2)
    doc.add_paragraph("按顺序执行：")
    _add_code_block(
        doc,
        "# ① 长者智能计算图 DataAgent\n"
        "python /srv/supercare/data_agent_task1.py \\\n"
        '  --input "/srv/supercare/DataSource/陈女士/ms_chen_CareCase.xlsx" \\\n'
        '  --output-root "/srv/supercare/data-agent1" \\\n'
        '  --config "/srv/supercare/config/data_agent1_config.json" \\\n'
        "  --parser-mode openpyxl\n\n"
        "# ② 三超循证知识基座 DataAgent\n"
        "python /srv/supercare/data_agent_task2.py \\\n"
        '  --input-graph "/srv/supercare/data-agent1/ms_chen_CareCase/results/ms_chen_CareCase_elder_health_computing_graph.json" \\\n'
        '  --output-root "/srv/supercare/data-agent2" \\\n'
        '  --config "/srv/supercare/config/data_agent1_config.json"\n\n'
        "# ③ 序贯分期与贝叶斯风险融合推断算法\n"
        "python /srv/supercare/task-agent/run_staging_pipeline.py\n"
        "python /srv/supercare/task-agent/generate_a0_visualizations.py\n\n"
        "# ④ 超级「GP-护士-照护员」智能体协同算法\n"
        "python /srv/supercare/task-agent/test_agent_a3_cda_hub.py\n"
        "# python /srv/supercare/task-agent/run_agents_a1_a14.py   # 全量 A1–A14",
    )

    # §3 成功验证
    doc.add_heading("3. 成功验证清单", level=1)
    _add_table(
        doc,
        ["模块", "关键产物"],
        [
            [
                "长者智能计算图 DataAgent",
                "data-agent1/.../ms_chen_CareCase_elder_health_computing_graph.json",
            ],
            ["三超循证知识基座 DataAgent", "data-agent2/batch_summary.json"],
            ["序贯分期与贝叶斯风险融合推断算法", "task-agent/output/a0_序贯疾病分期结果.json"],
            ["超级「GP-护士-照护员」智能体协同算法", "task-agent/output/a3_GP协作专业答复.pdf"],
        ],
    )
    _add_code_block(doc, f"curl -s {API_BASE}/outputs | python3 -m json.tool")

    # §4 产物说明
    doc.add_heading("4. 产物说明（按模块）", level=1)
    _add_output_catalog_tables(doc)

    # §5 常见问题
    doc.add_heading("5. 常见问题", level=1)
    _add_faq(
        doc,
        "qwen.api_key 未填写",
        "编辑 config/data_agent1_config.json 填入 DashScope Key 后重启 API。",
    )
    _add_faq(
        doc,
        "MinerU 正式模式报错",
        "确认 mineru.api_url 为可访问的 http/https 地址；或保持 fast_mode=true 使用 openpyxl。",
    )
    _add_faq(
        doc,
        "DataAgent2 报找不到图谱",
        "先完成 DataAgent1，确认 elder_health_computing_graph.json 已生成。",
    )
    _add_faq(
        doc,
        "中文 PDF/图片乱码",
        "安装中文字体 wqy-zenhei 或项目内 task-agent/fonts/ 思源黑体。",
    )

    # §6 相关文档
    doc.add_heading("6. 相关文档", level=1)
    _add_bullet(doc, "API 源码：/srv/supercare/supercare_quickstart_api.py")
    _add_bullet(doc, "Markdown 版：比赛文档/快速运行指南.md")
    _add_bullet(doc, "详细部署：比赛文档/3_系统部署与运行说明_超级[GP-护士-照护员]智能体协同算法.docx")
    _add_bullet(doc, "测试日志：比赛文档/4_系统运行日志与对比实验_超级[GP-护士-照护员]智能体协同算法.docx")

    return doc


def main() -> None:
    document = build_quickstart_guide()
    GUIDE_DOC.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(GUIDE_DOC))
    DELIVERABLE_DOC.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(DELIVERABLE_DOC))
    print(f"已生成: {GUIDE_DOC}")
    print(f"已同步: {DELIVERABLE_DOC}")


if __name__ == "__main__":
    main()
