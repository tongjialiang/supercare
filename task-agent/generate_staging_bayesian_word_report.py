#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《本地序贯分期与贝叶斯风险融合》算法说明 Word 文档。

- 面向算法原理、公式与文献先验，非单病例临床报告；
- 公式采用 Office Math（OMML），在 Word 中可继续用公式编辑器/MathType 美化。
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from docx_equations import (
    EQUATION_LOGISTIC_LEVEL,
    EQUATION_LOG_LIKELIHOOD,
    EQUATION_LR_Z_NEG,
    EQUATION_LR_Z_POS,
    EQUATION_ODDS,
    EQUATION_ODDS_UPDATE,
    EQUATION_POSTERIOR,
    EQUATION_PRIOR_STAGE,
    EQUATION_ZSCORE,
    add_display_equation,
)
from staging.prior_literature import (
    PRIOR_ASSUMPTION_H,
    PRIOR_FORMULA_NOTE,
    PRIOR_REFERENCES,
    STAGE_BPSD_PRIORS_LITERATURE,
    STAGE_LABELS,
    STAGE_PRIOR_DERIVATION,
)

PROJECT_ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = PROJECT_ROOT / "output"
COMPETITION_DELIVERABLE = Path("/srv/supercare/比赛文档/交付物")

STAGING_JSON = OUTPUT_DIR / "a0_序贯疾病分期结果.json"
BAYESIAN_JSON = OUTPUT_DIR / "a0_贝叶斯风险后验.json"
WORD_OUTPUT = OUTPUT_DIR / "本地序贯分期与贝叶斯风险融合_算法说明.docx"
DELIVERABLE_WORD = COMPETITION_DELIVERABLE / WORD_OUTPUT.name

# 配图为流水线运行示例（默认焦点病例），非本文档论述对象
EXAMPLE_FIGURES = [
    ("图1 疾病进展曲线示意（横轴=共同分期，纵轴=归一化进展指数）", "a0_疾病进展曲线_陈女士.png"),
    ("图2 序贯事件槽位与 MCMC 分期示意", "a0_序贯事件进展图_陈女士.png"),
    ("图3 贝叶斯先验—后验对比示意", "a0_贝叶斯先验后验图_陈女士.png"),
    ("图4 纵向序贯进展热图示意", "a0_生物标志物分期进展热图_陈女士.png"),
]


def _load_json(path: Path) -> Optional[dict]:
    if not path.is_file():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(6)


def _add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"• {text}")
    paragraph.paragraph_format.space_after = Pt(4)


def _add_picture(doc: Document, caption: str, image_path: Path) -> None:
    _add_paragraph(doc, caption, bold=True)
    if image_path.is_file():
        doc.add_picture(str(image_path), width=Inches(5.8))
    else:
        _add_paragraph(doc, f"（配图未生成：{image_path.name}，请先运行 generate_a0_visualizations.py）")
    doc.add_paragraph()


def _add_equation_block(doc: Document, mathml: str, caption: str = "") -> None:
    add_display_equation(doc, mathml, caption=caption)
    doc.add_paragraph()


def build_algorithm_document(
    staging: Optional[dict] = None,
    bayesian: Optional[dict] = None,
) -> Document:
    doc = Document()

    title = doc.add_heading(
        "SuperCare A0：本地序贯分期与贝叶斯风险融合\n算法说明",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_paragraph(
        doc,
        "文档性质：算法与计算方法说明书（非个体病例临床报告）。\n"
        "适用对象：评审专家、算法审计、超级 GP/护士智能体协同研发。\n"
        "运行环境：本地 DataSource → run_staging_pipeline.py，无外部分期 API。",
    )

    # 一、总览
    _add_heading(doc, "一、算法总览与数据流", level=1)
    _add_paragraph(
        doc,
        "A0 流水线在机构队列（默认 10 例）上完成：① 生物标志物提取；② 序贯事件异常模型 + MCMC 共同分期；"
        "③ 文献校准贝叶斯先验 + 标志物似然更新后验；④ 输出 JSON/配图供 A3/A8/A15 工具调用。",
    )
    _add_bullet(doc, "输入：DataSource 标准化 Excel（NPI、MMSE/MoCA 认知逆指标、Barthel/ADL；不含血压/血氧参与分期）")
    _add_bullet(doc, "输出：a0_序贯疾病分期结果.json、a0_贝叶斯风险后验.json、进展曲线与先验后验配图")
    _add_bullet(doc, "分期范式：n 个 BPSD 标志物 → n 条疾病进展曲线；横轴为共同分期 1–5，纵轴为标志物水平")

    # 二、序贯事件模型
    _add_heading(doc, "二、序贯事件异常模型", level=1)
    _add_paragraph(
        doc,
        "每个生物标志物 k 对应一个序贯「事件」：在共同疾病分期轴 s∈{1,…,5} 上，用 logistic 函数描述"
        "从基线水平 b_k 向异常水平 a_k 的过渡；事件 onset 为 τ_k，宽度为 w_k。",
    )
    _add_equation_block(doc, EQUATION_LOGISTIC_LEVEL, "式(1) 标志物 k 在分期 s 的期望水平")
    _add_paragraph(doc, "符号说明：")
    _add_bullet(doc, "b_k、a_k：由队列该标志物分位数（P25/P75）标定")
    _add_bullet(doc, "τ_k：随亚型内事件顺序后移，体现「认知→行为→功能」等序贯")
    _add_bullet(doc, "亚型：三种预设事件顺序；MCMC 在亚型间切换以拟合观测")

    _add_heading(doc, "2.1 MCMC 似然与推断", level=2)
    _add_equation_block(doc, EQUATION_LOG_LIKELIHOOD, "式(2) 给定分期 s 与亚型的对数似然（独立高斯残差近似）")
    _add_paragraph(
        doc,
        "采用 Metropolis-Hastings：对每位病例交替提议亚型与分期 s，以接受率监控混合；"
        "burn-in 后取分期后验中位数作为 disease_stage。",
    )

    # 三、进展曲线
    _add_heading(doc, "三、疾病进展曲线（可视化语义）", level=1)
    _add_paragraph(
        doc,
        "对每个亚型，在分期 1–5 上计算式(1) 的拟合值（fitted_levels），并按 MCMC 分期分桶求队列经验均值；"
        "多标志物绘于同一图时需归一化进展指数 0–100% 以便比较。",
    )
    _add_picture(doc, EXAMPLE_FIGURES[0][0], OUTPUT_DIR / EXAMPLE_FIGURES[0][1])
    _add_picture(doc, EXAMPLE_FIGURES[1][0], OUTPUT_DIR / EXAMPLE_FIGURES[1][1])

    # 四、贝叶斯
    _add_heading(doc, "四、贝叶斯风险融合", level=1)
    _add_heading(doc, "4.1 假设与先验（文献校准）", level=2)
    _add_paragraph(doc, PRIOR_ASSUMPTION_H)
    _add_paragraph(doc, PRIOR_FORMULA_NOTE)
    _add_equation_block(doc, EQUATION_PRIOR_STAGE, "式(3) 分期先验（文献患病率 × 月尺度恶化系数）")

    prior_table = doc.add_table(rows=1, cols=3)
    prior_table.style = "Table Grid"
    prior_table.rows[0].cells[0].text = "分期 s"
    prior_table.rows[0].cells[1].text = "P(H|s)"
    prior_table.rows[0].cells[2].text = "文献推导摘要"
    for stage_index in range(1, 6):
        row = prior_table.add_row().cells
        row[0].text = f"{stage_index} — {STAGE_LABELS.get(stage_index, '')}"
        row[1].text = f"{STAGE_BPSD_PRIORS_LITERATURE[stage_index]:.0%}"
        row[2].text = STAGE_PRIOR_DERIVATION.get(stage_index, "")
    doc.add_paragraph()

    _add_heading(doc, "4.2 先验概率参考文献", level=2)
    for reference in PRIOR_REFERENCES:
        _add_paragraph(
            doc,
            f"[{reference['id']}] {reference['citation']}\n    用途：{reference['use']}",
        )

    _add_heading(doc, "4.3 似然更新与后验", level=2)
    _add_equation_block(doc, EQUATION_ZSCORE, "式(4) 队列 z-score 标准化")
    _add_equation_block(doc, EQUATION_LR_Z_POS, "式(5a) z_i > 0 时似然比")
    _add_equation_block(doc, EQUATION_LR_Z_NEG, "式(5b) z_i < 0 时似然比（取倒数）")
    _add_equation_block(doc, EQUATION_ODDS, "式(6) 先验 odds")
    _add_equation_block(doc, EQUATION_ODDS_UPDATE, "式(7) odds 更新")
    _add_equation_block(doc, EQUATION_POSTERIOR, "式(8) 后验概率")

    _add_paragraph(
        doc,
        "决策阈值（与超级 GP 协同）：后验 P(H|D) ≥ 35% 建议会诊；≥ 25% 护士周评 NPI；≥ 20% 照护员加强监测。",
    )
    _add_picture(doc, EXAMPLE_FIGURES[2][0], OUTPUT_DIR / EXAMPLE_FIGURES[2][1])

    # 五、实现
    _add_heading(doc, "五、本地实现与模块", level=1)
    _add_bullet(doc, "staging/biomarker_extraction.py — DataSource → 队列矩阵")
    _add_bullet(doc, "staging/sustain_mcmc_model.py — 序贯事件模型 + MCMC + 进展曲线")
    _add_bullet(doc, "staging/prior_literature.py — 文献先验表与参考文献")
    _add_bullet(doc, "staging/bayesian_risk.py — 先验 + z-score 似然 + 后验")
    _add_bullet(doc, "generate_a0_visualizations.py — 配图；common_utils 工具注入 A3/A8/A15")

    _add_heading(doc, "六、运行示例配图", level=1)
    _add_paragraph(
        doc,
        "下列配图来自默认流水线的一次运行结果（焦点病例仅作算例，不代表本文档主体）。"
        "换队列或焦点病例后，图形数值会变，算法结构不变。",
    )
    _add_picture(doc, EXAMPLE_FIGURES[3][0], OUTPUT_DIR / EXAMPLE_FIGURES[3][1])

    # 附录算例
    if staging and bayesian:
        _add_heading(doc, "附录 A：一次运行的数值算例（仅供核对公式）", level=1)
        focus = staging.get("focus_case_staging", {})
        _add_paragraph(
            doc,
            f"算例病例 ID：{focus.get('case_id', '—')}（{focus.get('folder_name', '—')}）；"
            f"MCMC 阶段 {focus.get('disease_stage')}；文献先验 "
            f"{bayesian.get('priors', {}).get('bpsd_escalation_30d', 0):.1%} → 后验 "
            f"{bayesian.get('posteriors', {}).get('bpsd_escalation_30d', 0):.1%}。",
        )
        calc = bayesian.get("calculation_steps", {})
        _add_bullet(doc, f"先验 odds = {calc.get('prior_odds')}")
        _add_bullet(doc, f"合并 LR = {bayesian.get('combined_likelihood_ratio')}")
        _add_bullet(doc, f"后验 odds = {calc.get('posterior_odds')}")

    _add_heading(doc, "附录 B：本地命令", level=1)
    _add_bullet(doc, "python run_staging_pipeline.py")
    _add_bullet(doc, "python generate_a0_visualizations.py")
    _add_bullet(doc, "python generate_staging_bayesian_word_report.py")

    _add_paragraph(
        doc,
        "提示：在 Microsoft Word 中打开后，双击公式可进入「公式工具」进一步调整为专业型（类似 MathType 排版）。",
        bold=True,
    )

    return doc


def main() -> None:
    staging = _load_json(STAGING_JSON)
    bayesian = _load_json(BAYESIAN_JSON)
    doc = build_algorithm_document(staging, bayesian)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(WORD_OUTPUT))
    COMPETITION_DELIVERABLE.mkdir(parents=True, exist_ok=True)
    doc.save(str(DELIVERABLE_WORD))
    print(f"已生成算法说明 Word：\n  {WORD_OUTPUT}\n  {DELIVERABLE_WORD}")


if __name__ == "__main__":
    main()
