#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将序贯分期与贝叶斯风险融合内容深度写入技术报告（非段末追加），
统一学术表述，避免 DataSource / tool_ 等工程命名。
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Callable, List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

COMPETITION = Path("/srv/supercare/比赛文档")
SOURCE_REPORT = COMPETITION / "1_技术报告_超级[GP-护士-照护员]智能体协同算法.docx"
OUTPUT_REPORT = (
    COMPETITION / "1_技术报告_超级[GP-护士-照护员]智能体协同算法_含序贯贝叶斯_20260525.docx"
)
DELIVERABLE = COMPETITION / "交付物" / OUTPUT_REPORT.name

# 需删除的“段末追加”或重复块（全文子串匹配即删）
PARAGRAPHS_TO_REMOVE_SUBSTRINGS = [
    "【A0 本地计算层】",
    "A0 本地序贯分期与贝叶斯层：DataSource",
    "输出：a0_序贯疾病分期结果.json、a0_贝叶斯风险后验.json",
    "在 11 类标准化图谱工具基础上，新增 tool_",
    " 另含《含序贯贝叶斯》修订版及 A0 JSON/配图",
]

# 独立标题段（整段等于下列文本时删除，避免创新点重复出现在 DataAgent 前）
HEADINGS_TO_REMOVE_EXACT = {
    "A0 本地序贯分期与贝叶斯风险融合",
    "A0 本地序贯分期与贝叶斯工具",
}

ABSTRACT_INTRO_OLD = "以“数据—知识—应用”三位一体为核心的全链路智能体解决方案"
ABSTRACT_INTRO_NEW = (
    "以「数据—知识—推断—协同」四层贯通架构为核心的全链路智能体解决方案"
)

ABSTRACT_INFERENCE_PARAGRAPH = (
    "再次，提出序贯分期与贝叶斯风险融合推断算法，作为四层架构中「推断层」的核心方法："
    "依托爱照护专业照护机构长期积累的纵向评估数据，提取神经精神症状量表（NPI）、"
    "认知筛查（MMSE/MoCA）与日常生活能力（Barthel）等 BPSD 相关指标；"
    "将各指标建模为随病程逐步由正常转向异常的序贯过程，用 logistic 曲线刻画异常出现的时机与幅度，"
    "并统一映射到 1—5 级疾病分期标尺，通过马尔可夫链蒙特卡洛（MCMC）估计长者当前所处阶段；"
    "在分期基础上，按病程阶段引用国内外文献给出的分层先验概率 P(H|stage)，"
    "再将该长者各指标相对爱照护照护队列平均水平的高低（标准化 z 分数）纳入贝叶斯更新，"
    "经似然比递推得到未来 30 天内 BPSD 临床显著恶化的后验概率；"
    "输出疾病分期洞察与风险后验等结果，在返院 D2、D7、D30 等节点以 35%、25%、20% "
    "的阈值支撑升级与会诊决策，使长期照护「证据轮」获得可解释、可复现、可校准的量化风险表达。"
)

ABSTRACT_SUPER_GP_PARAGRAPH = (
    "最后，在前述数据底座、循证知识与量化推断支撑下，构建超级[GP—护士—照护员]智能体协同体系，"
    "形成「1 个总控智能体 + 3 个角色中枢 + 11 个任务型智能体 + 十三类标准化工具」的协同架构，"
    "将分期洞察、风险后验与图谱洞察工具贯通至 GP、护士与照护员决策链，"
    "实现养老照护全流程的数据驱动、知识赋能、决策智能与执行规范深度融合。"
    "A/B 测试显示，启用工具链后智能体输出完整度提升 23.04%，"
    "照护方案的结构化程度与可执行性显著增强。"
)

ABSTRACT_BPSD_PARAGRAPH = (
    "以 BPSD 患者返院全周期为典型应用场景，将照护过程划分为 D-2、D0、D2、D5、D7、D14、D30 "
    "七个关键阶段，在序贯分期与贝叶斯后验推断及多智能体协同、人在回路机制共同支撑下，"
    "形成「数据—知识—推断—决策—执行—反馈」闭环，并进一步沉淀照护科学实践图谱与高价值行业语料库，"
    "为养老照护知识标准化、经验传承和规模化应用提供技术路径。"
)

PROJECT_GOAL_INFERENCE = (
    "序贯分期与贝叶斯风险融合推断引擎：依托爱照护纵向评估数据，完成生物标志物提取、"
    "序贯分期推断、文献校准先验与个体似然融合，输出可解释的疾病分期、进展曲线与升级后验，"
    "为超级 GP 协同层提供量化风险底座。"
)

INNOVATION_TITLE = "序贯分期与文献校准贝叶斯风险融合技术"
INNOVATION_BODY = (
    "创新性构建多标志物序贯分期模型：各指标按 logistic 曲线描述由正常至异常的变化过程，"
    "共用五级分期标尺，由 MCMC 估计当前疾病阶段；在此基础上构建贝叶斯风险层，"
    "以文献分层先验 P(H|stage) 刻画分期相关的基线恶化概率，"
    "再结合个体指标相对爱照护队列的 z 分数经似然比更新后验，"
    "并与 GP 会诊阈值联动，实现返院全周期量化预警与分级决策。"
)

FUNCTION_GOAL_SUFFIX = (
    "体系在健康计算图与循证知识基座之上，增设序贯分期与贝叶斯风险融合推断层，"
    "将 NPI、认知量表、Barthel 等纵向评估转化为共同分期与升级后验，"
    "为事件分级、会诊准备与再稳定评估提供可计算、可解释的风险依据。"
)

ARCHITECTURE_LAYER = (
    "协同架构在「健康计算图—循证知识—多智能体执行」主链中嵌入序贯—贝叶斯量化推断层："
    "纵向评估队列经序贯事件建模与 MCMC 分期后，进入文献校准贝叶斯更新，"
    "分期洞察与风险后验作为扩展工具供总控与各角色中枢在 D2/D7/D30 等节点调用，"
    "与升级判断、会诊支撑等任务型智能体形成「量化风险—规则证据—角色判断」三层递进。"
    "算法原理、公式推导与文献先验校准详见「五、面向 BPSD 返院照护的序贯分期与贝叶斯风险融合算法」。"
)

TOOLS_INTRO = (
    "在十一类健康计算图洞察工具基础上，扩展疾病分期洞察与风险后验两类量化推断工具，"
    "共十三类标准化工具，为智能体决策提供结构化证据与可计算风险支撑："
)

TOOL_ROW_STAGING = (
    "疾病分期洞察工具",
    "输出 MCMC 疾病分期、序贯事件参数与多标志物进展曲线，支撑病程定位与趋势研判",
)
TOOL_ROW_BAYESIAN = (
    "风险后验工具",
    "提供文献先验、标志物似然更新后的升级后验及会诊阈值对照，支撑 GP 介入与事件升级决策",
)

WORKFLOW_PATCHES = {
    "D2 事件升级与适应观察阶段": (
        "核心目标：对 BPSD 事件进行结构化记录、分级和升级建议生成；"
        "可结合序贯分期结果与升级后验概率，为观察/护士介入/GP 介入提供量化依据"
    ),
    "D7 会诊准备与协同阶段": (
        "核心目标：形成一周综合判断材料，为 GP/精神科协作与会诊提供依据；"
        "综合报告纳入贝叶斯后验、共同分期与三时点功能对比，支撑会诊阈值判断"
    ),
    "D30 再稳定评估阶段": (
        "核心目标：形成结案建议、再稳定判断与高价值训练语料沉淀；"
        "对照返院初期分期/后验与当前评估，评价干预成效与再稳定程度"
    ),
}

CHAPTER_TITLE_OLD = "A0 本地序贯分期与贝叶斯风险融合"
CHAPTER_TITLE_NEW = "序贯分期与贝叶斯风险融合"
CHAPTER_INTRO = (
    "本节阐述序贯分期与贝叶斯风险融合推断算法的原理、"
    "核心公式、文献先验来源与示意配图；配图为流水线运行示例，不代表单一病例临床报告。"
)

FIGURE_CAPTION_REPLACEMENTS = {
    "图A0-1 整体叙事：DataSource → 序贯分期 → 贝叶斯 → 超级GP": (
        "图A0-1 整体技术链路：纵向评估队列 → 序贯分期推断 → 贝叶斯风险融合 → 超级GP协同决策"
    ),
}

CONCLUSION_OLD_SNIPPET = (
    "构建了由长者健康计算图DataAgent、三超循证知识基座DataAgent和超级[GP-护士-照护员]智能体协同算法组成的全链路智能体系。"
)
CONCLUSION_NEW_SNIPPET = (
    "构建了由长者健康计算图 DataAgent、三超循证知识基座 DataAgent、"
    "序贯分期与贝叶斯风险融合推断算法以及超级[GP-护士-照护员]智能体协同算法组成的四层全链路智能体系。"
)

CONCLUSION_TAIL = (
    "序贯分期与文献校准贝叶斯更新将返院高风险判读从经验阈值推进为可解释的后验概率表达，"
    "与四层语料加工、多源质控体系形成互补。"
)

# 全文遗留表述批量替换（专章、创新点等）
DOCUMENT_WORDING_REPLACEMENTS = [
    (
        "创新性提出「序贯共同分期」范式：每个生物标志物对应一个序贯事件（logistic onset），"
        "全体标志物共用分期轴 1–5，由 MCMC 推断亚型与阶段；贝叶斯层采用国内外文献分层先验 "
        "P(H|stage)≈p_s×λ_s，再经标志物似然比更新后验，驱动 GP 会诊阈值（35%/25%/20%）。",
        INNOVATION_BODY,
    ),
    (
        "「序贯」含义：多个 BPSD 生物标志物按预设顺序依次出现异常，全体标志物处在同一条疾病分期轴 "
        "s∈{1,…,5} 上，而非各指标单独分期再平均。",
        "「序贯」含义：各 BPSD 相关指标随病程逐步由正常转向异常，并在同一套 1—5 级分期标尺上刻画当前所处阶段，"
        "避免对各指标分别分期后再简单平均。",
    ),
    (
        "每个标志物 k 对应序贯事件模型：baseline b_k、异常水平 a_k、onset τ_k；"
        "分期越靠后的事件 onset 越晚。亚型由事件顺序定义（如认知→MoCA→NPI→ADL 等三种轨迹）。",
        "每个标志物 k 用序贯事件模型描述：正常基线 b_k、异常水平 a_k、异常出现时刻 τ_k；"
        "分期越靠后，各指标异常出现的时间通常越晚。",
    ),
    (
        "Metropolis-Hastings MCMC 在队列最新截面上联合推断亚型与分期；",
        "Metropolis-Hastings MCMC 在爱照护队列最新评估截面上推断当前疾病分期；",
    ),
    (
        "贝叶斯层假设 H：30 天内 BPSD 临床显著恶化（NPI 总分升高≥4 或任一分项≥4）。"
        "先验 P(H) 由 MCMC 分期查文献校准表；各标志物队列 z-score 映射为连续似然比 LR，"
        "odds 更新后得到后验，与 GP 阈值联动。",
        "贝叶斯层关注事件 H：未来 30 天内 BPSD 是否临床显著恶化（NPI 总分升高≥4 或任一分项≥4）。"
        "先验 P(H) 由 MCMC 所得分期对照文献校准表确定；各标志物相对爱照护队列的 z 分数"
        "转换为似然比后更新 odds，得到后验概率，并与 GP 会诊阈值联动。",
    ),
    ("机构纵向照护评估队列", "爱照护专业照护机构纵向评估数据"),
    ("同龄队列", "爱照护照护队列"),
    ("恶化顺序亚型", "疾病分期"),
    ("恶化亚型", "疾病分期"),
    ("推断亚型与分期", "推断疾病分期"),
    ("推断亚型与阶段", "推断疾病分期"),
    ("亚型与分期", "疾病分期"),
    ("分期亚型", "疾病分期"),
    ("序贯共同分期", "序贯分期"),
]


def _paragraph_text(paragraph: Paragraph) -> str:
    return (paragraph.text or "").strip()


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def _find_paragraph(
    doc: Document,
    keyword: str,
    *,
    style_contains: str = "",
    exact: bool = False,
) -> Optional[Paragraph]:
    for paragraph in doc.paragraphs:
        text = _paragraph_text(paragraph)
        if exact:
            if text != keyword:
                continue
        elif keyword not in text:
            continue
        if style_contains and style_contains not in (paragraph.style.name or ""):
            continue
        return paragraph
    return None


def _remove_appendix_paragraphs(doc: Document) -> None:
    to_remove: List[Paragraph] = []
    for paragraph in doc.paragraphs:
        text = _paragraph_text(paragraph)
        if text in HEADINGS_TO_REMOVE_EXACT:
            to_remove.append(paragraph)
            continue
        if any(substr in text for substr in PARAGRAPHS_TO_REMOVE_SUBSTRINGS):
            to_remove.append(paragraph)
            continue
        if text.startswith("创新性提出「序贯共同分期」范式") and "tool_" not in text:
            # 误插在 DataAgent 前的创新点正文副本
            to_remove.append(paragraph)
    for paragraph in to_remove:
        _remove_paragraph(paragraph)


def _insert_after(anchor: Paragraph, text: str, style: str = "Body Text") -> Paragraph:
    new_element = OxmlElement("w:p")
    anchor._element.addnext(new_element)
    new_paragraph = Paragraph(new_element, anchor._parent)
    new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def _patch_abstract_intro(doc: Document) -> None:
    """开篇由「三位一体」调整为「四层贯通」。"""
    paragraph = _find_paragraph(doc, ABSTRACT_INTRO_OLD)
    if paragraph:
        text = _paragraph_text(paragraph).replace(ABSTRACT_INTRO_OLD, ABSTRACT_INTRO_NEW, 1)
        _replace_paragraph_text(paragraph, text)


def _patch_abstract_tool_count(doc: Document) -> None:
    """摘要中协同架构的工具数量与正文十三类保持一致。"""
    paragraph = _find_paragraph(doc, "11 个任务型智能体 + 11 类标准化工具")
    if paragraph:
        text = _paragraph_text(paragraph)
        text = text.replace(
            "11 个任务型智能体 + 11 类标准化工具",
            "11 个任务型智能体 + 十三类标准化工具",
        )
        _replace_paragraph_text(paragraph, text)


def _patch_abstract(doc: Document) -> None:
    """
    摘要四项顺序：首先（数据）→ 其次（知识）→ 再次（序贯—贝叶斯，重点）→ 最后（协同，收束）。
    「最后」之后仅保留应用场景总述，不再追加新算法内容。
    """
    inference_paragraph = None
    for keyword in (
        "再次，提出序贯分期与贝叶斯风险融合推断算法",
        "此外，提出序贯分期与贝叶斯风险融合推断算法",
        "第三，提出序贯分期",
        "此外，面向 BPSD 返院场景",
    ):
        inference_paragraph = _find_paragraph(doc, keyword)
        if inference_paragraph:
            break

    super_gp_paragraph = _find_paragraph(
        doc, "最后，提出基于超级[GP—护士—照护员]智能体协同算法"
    )
    if not super_gp_paragraph:
        super_gp_paragraph = _find_paragraph(
            doc, "最后，在前述数据底座、循证知识与量化推断支撑下"
        )

    if inference_paragraph:
        _replace_paragraph_text(inference_paragraph, ABSTRACT_INFERENCE_PARAGRAPH)
    elif super_gp_paragraph:
        # 在超级 GP 段之前插入推断段
        new_element = OxmlElement("w:p")
        super_gp_paragraph._element.addprevious(new_element)
        inference_paragraph = Paragraph(new_element, super_gp_paragraph._parent)
        inference_paragraph.style = "Body Text"
        inference_paragraph.add_run(ABSTRACT_INFERENCE_PARAGRAPH)

    if super_gp_paragraph:
        _replace_paragraph_text(super_gp_paragraph, ABSTRACT_SUPER_GP_PARAGRAPH)

    # 保证阅读顺序：再次段在最后段之前（若历史上顺序颠倒则交换相邻段落）
    if inference_paragraph and super_gp_paragraph:
        inference_element = inference_paragraph._element
        super_gp_element = super_gp_paragraph._element
        if list(inference_element.getparent()).index(inference_element) > list(
            super_gp_element.getparent()
        ).index(super_gp_element):
            super_gp_element.addprevious(inference_element)

    bpsd_paragraph = _find_paragraph(doc, "以 BPSD 患者返院全周期为典型应用场景")
    if bpsd_paragraph:
        _replace_paragraph_text(bpsd_paragraph, ABSTRACT_BPSD_PARAGRAPH)


def _patch_project_goals(doc: Document) -> None:
    goals_anchor = _find_paragraph(
        doc,
        "超级[GP—护士—照护员]智能体协同算法：面向业务协同与自动执行任务",
    )
    goal_line = _find_paragraph(doc, "序贯分期与贝叶斯风险融合推断引擎")
    if not goal_line:
        goal_line = _find_paragraph(doc, "序贯分期与贝叶斯风险融合引擎")
    if goal_line:
        _replace_paragraph_text(goal_line, PROJECT_GOAL_INFERENCE)
    elif goals_anchor:
        _insert_after(goals_anchor, PROJECT_GOAL_INFERENCE)


def _patch_innovation_section(doc: Document) -> None:
    workflow_innovation = _find_paragraph(
        doc, "通过多智能体协同技术自动从全周期照护报告中提取可迁移的照护方法论"
    )
    if not workflow_innovation:
        return

    old_titles = (
        "序贯共同分期与文献校准贝叶斯风险融合技术",
        INNOVATION_TITLE,
    )
    for title in old_titles:
        heading = _find_paragraph(doc, title, exact=True)
        if heading:
            _replace_paragraph_text(heading, INNOVATION_TITLE)
            next_el = heading._element.getnext()
            if next_el is not None:
                body = Paragraph(next_el, heading._parent)
                body_text = _paragraph_text(body)
                if body_text.startswith("创新性"):
                    _replace_paragraph_text(body, INNOVATION_BODY)
                    return
            _insert_after(heading, INNOVATION_BODY)
            return

    heading = _insert_after(workflow_innovation, INNOVATION_TITLE, "Heading 3")
    _insert_after(heading, INNOVATION_BODY)


def _patch_super_gp_function_goal(doc: Document) -> None:
    goal = _find_paragraph(
        doc,
        "基于\"基态-触发-就医-返回\"全周期照护闭环设计理念，构建覆盖BPSD患者返院全周期的智能照护体系",
    )
    if not goal:
        return
    text = _paragraph_text(goal)
    if "序贯分期与贝叶斯风险融合推断层" in text:
        return
    if "本地序贯分期" in text:
        text = text.replace("本地序贯分期与贝叶斯风险融合推断层", "序贯分期与贝叶斯风险融合推断层")
        _replace_paragraph_text(goal, text)
        return
    _replace_paragraph_text(goal, text + FUNCTION_GOAL_SUFFIX)


def _patch_architecture(doc: Document) -> None:
    existing = _find_paragraph(doc, "协同架构在「健康计算图—循证知识—多智能体执行」")
    if existing:
        _replace_paragraph_text(existing, ARCHITECTURE_LAYER)
        return
    arch_intro = _find_paragraph(doc, "采用\"1个总控智能体+3个角色中枢+11个任务型智能体\"的架构设计")
    if arch_intro:
        _insert_after(arch_intro, ARCHITECTURE_LAYER)


def _patch_tools_table(doc: Document) -> None:
    tools_intro = _find_paragraph(doc, "设计了11类基于健康计算图的工具")
    if tools_intro:
        _replace_paragraph_text(tools_intro, TOOLS_INTRO)

    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header = table.rows[0].cells[0].text.strip()
        if header != "工具名称":
            continue
        names = [row.cells[0].text.strip() for row in table.rows[1:]]
        if "疾病分期洞察工具" in names:
            for row in table.rows[1:]:
                if row.cells[0].text.strip() == "疾病分期洞察工具":
                    row.cells[1].text = TOOL_ROW_STAGING[1]
            return
        row_staging = table.add_row()
        row_staging.cells[0].text = TOOL_ROW_STAGING[0]
        row_staging.cells[1].text = TOOL_ROW_STAGING[1]
        row_bayes = table.add_row()
        row_bayes.cells[0].text = TOOL_ROW_BAYESIAN[0]
        row_bayes.cells[1].text = TOOL_ROW_BAYESIAN[1]
        return


def _patch_workflow_stages(doc: Document) -> None:
    for keyword, new_core in WORKFLOW_PATCHES.items():
        paragraph = _find_paragraph(doc, keyword)
        if not paragraph:
            continue
        next_paragraph = paragraph._element.getnext()
        if next_paragraph is None:
            continue
        sibling = Paragraph(next_paragraph, paragraph._parent)
        if _paragraph_text(sibling).startswith("核心目标："):
            _replace_paragraph_text(sibling, new_core)


def _patch_dedicated_chapter(doc: Document) -> None:
    chapter_heading = _find_paragraph(doc, CHAPTER_TITLE_OLD, exact=True)
    if chapter_heading:
        _replace_paragraph_text(chapter_heading, CHAPTER_TITLE_NEW)

    chapter_intro = _find_paragraph(doc, "本章说明 A0 本地序贯疾病分期")
    if chapter_intro:
        _replace_paragraph_text(chapter_intro, CHAPTER_INTRO)

    for paragraph in doc.paragraphs:
        text = _paragraph_text(paragraph)
        for old_caption, new_caption in FIGURE_CAPTION_REPLACEMENTS.items():
            if old_caption in text:
                _replace_paragraph_text(paragraph, new_caption)
        if text == "数据流与模块":
            _replace_paragraph_text(paragraph, "处理流程与模块")
        if "staging/biomarker_extraction.py" in text:
            _replace_paragraph_text(
                paragraph,
                "处理链路依次为：生物标志物提取 → 序贯事件与 MCMC 分期 → 文献先验构建 → 贝叶斯后验更新。",
            )
        if text.startswith("产出：a0_"):
            _replace_paragraph_text(
                paragraph,
                "输出包括序贯分期结构化结果、风险后验结果及疾病进展与先验—后验对比配图，供协同智能体调用。",
            )
        if "本地命令：python task-agent" in text or "机构本地环境" in text:
            _replace_paragraph_text(
                paragraph,
                "推断流水线在照护机构侧部署运行，算法细节见《序贯分期与贝叶斯风险融合推断算法说明》专册。",
            )


def _patch_conclusion(doc: Document) -> None:
    conclusion = _find_paragraph(doc, CONCLUSION_OLD_SNIPPET[:40])
    if not conclusion:
        return
    text = _paragraph_text(conclusion)
    text = text.replace("本地序贯分期与贝叶斯风险融合推断层", "序贯分期与贝叶斯风险融合推断算法")
    if "序贯分期与贝叶斯风险融合推断算法" in text and "四层全链路" in text:
        if CONCLUSION_TAIL.strip() not in text:
            text += CONCLUSION_TAIL
        _replace_paragraph_text(conclusion, text)
        return
    text = text.replace(CONCLUSION_OLD_SNIPPET, CONCLUSION_NEW_SNIPPET)
    if not text.endswith("。"):
        text += "。"
    text += CONCLUSION_TAIL
    _replace_paragraph_text(conclusion, text)


def _patch_document_wording(doc: Document) -> None:
    """专章及正文遗留术语统一为通俗表述。"""
    for paragraph in doc.paragraphs:
        text = _paragraph_text(paragraph)
        if not text:
            continue
        new_text = text
        for old, new in DOCUMENT_WORDING_REPLACEMENTS:
            if old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != text:
            _replace_paragraph_text(paragraph, new_text)


def _strip_local_wording(doc: Document) -> None:
    """全文去除「本地」修饰，统一为算法/推断层表述。"""
    replacements = [
        ("本地序贯分期与贝叶斯风险融合推断方法", "序贯分期与贝叶斯风险融合推断算法"),
        ("本地序贯分期与贝叶斯风险融合推断算法", "序贯分期与贝叶斯风险融合推断算法"),
        ("本地序贯分期与贝叶斯风险融合", "序贯分期与贝叶斯风险融合"),
        ("本地序贯疾病分期与贝叶斯风险融合", "序贯分期与贝叶斯风险融合"),
        ("本地量化推断层", "序贯—贝叶斯量化推断层"),
        ("A0 本地序贯分期与贝叶斯风险融合", CHAPTER_TITLE_NEW),
        ("此外，提出序贯分期与贝叶斯风险融合推断算法", ABSTRACT_INFERENCE_PARAGRAPH),
        ("第三，提出序贯分期", ABSTRACT_INFERENCE_PARAGRAPH),
        ("第四，提出序贯分期", ABSTRACT_INFERENCE_PARAGRAPH),
    ]
    for paragraph in doc.paragraphs:
        text = _paragraph_text(paragraph)
        if not text:
            continue
        new_text = text
        for old, new in replacements:
            new_text = new_text.replace(old, new)
        if new_text != text:
            _replace_paragraph_text(paragraph, new_text)


def integrate_report(doc_path: Path) -> None:
    document = Document(str(doc_path))
    _remove_appendix_paragraphs(document)
    _patch_abstract_intro(document)
    _patch_abstract(document)
    _patch_abstract_tool_count(document)
    _patch_project_goals(document)
    _patch_innovation_section(document)
    _patch_super_gp_function_goal(document)
    _patch_architecture(document)
    _patch_tools_table(document)
    _patch_workflow_stages(document)
    _patch_dedicated_chapter(document)
    _patch_conclusion(document)
    _patch_document_wording(document)
    _strip_local_wording(document)
    document.save(str(doc_path))


def refine_existing_report() -> None:
    """在已生成的报告上微调，无需从原版重建。"""
    if not OUTPUT_REPORT.is_file():
        raise FileNotFoundError(f"未找到报告: {OUTPUT_REPORT}")
    integrate_report(OUTPUT_REPORT)
    DELIVERABLE.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUTPUT_REPORT, DELIVERABLE)
    print(f"已微调并保存: {OUTPUT_REPORT}")
    print(f"已同步交付物: {DELIVERABLE}")


def main() -> None:
    # 从原版重新生成含公式专章的完整稿，再执行深度融合（避免在已污染稿上重复修补）
    from update_competition_report_from_original import (  # noqa: WPS433
        _insert_a0_chapter_before_delivery,
        _patch_delivery_section,
        _patch_summary_and_innovation,
    )

    if not SOURCE_REPORT.is_file():
        raise FileNotFoundError(f"未找到原版: {SOURCE_REPORT}")

    shutil.copy2(SOURCE_REPORT, OUTPUT_REPORT)
    document = Document(str(OUTPUT_REPORT))
    _patch_summary_and_innovation(document)
    _insert_a0_chapter_before_delivery(document)
    _patch_delivery_section(document)
    document.save(str(OUTPUT_REPORT))

    integrate_report(OUTPUT_REPORT)

    DELIVERABLE.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUTPUT_REPORT, DELIVERABLE)
    print(f"已深度融合并保存: {OUTPUT_REPORT}")
    print(f"已同步交付物: {DELIVERABLE}")


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == "--refine":
        refine_existing_report()
    else:
        main()
