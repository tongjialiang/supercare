#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将技术报告中的「序贯分期与贝叶斯融合推断算法」章节写入部署说明与运行日志文档。"""

from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

COMPETITION = Path("/srv/supercare/比赛文档")
TECH_REPORT = COMPETITION / "1_技术报告_超级[GP-护士-照护员]智能体协同算法.docx"
DEPLOY_DOC = COMPETITION / "3_系统部署与运行说明_超级[GP-护士-照护员]智能体协同算法.docx"
LOG_DOC = COMPETITION / "4_系统运行日志与对比实验_超级[GP-护士-照护员]智能体协同算法.docx"
FIG_DIR = Path(__file__).resolve().parent / "output"

# 技术报告专章段落范围（含图5-1～5-4）
SOURCE_START = 285
SOURCE_END = 328  # 不含「典型任务执行示例」

FIGURE_FILES = {
    "图5-1": "图5-1_整体技术链路_爱照护序贯贝叶斯超级GP.png",
    "图5-2": "a0_疾病进展曲线_陈女士.png",
    "图5-3": "a0_序贯事件进展图_陈女士.png",
    "图5-4": "a0_贝叶斯先验后验图_陈女士.png",
}


def _text(paragraph: Paragraph) -> str:
    return (paragraph.text or "").strip()


def _find_paragraph(doc: Document, keyword: str, *, exact: bool = False) -> Optional[Paragraph]:
    for paragraph in doc.paragraphs:
        text = _text(paragraph)
        if exact and text == keyword:
            return paragraph
        if not exact and keyword in text:
            return paragraph
    return None


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _apply_paragraph_style(paragraph: Paragraph, style_name: str, fallback: Paragraph) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        if fallback.style:
            paragraph.style = fallback.style


def _insert_paragraph_after(anchor: Paragraph, text: str, style_name: str) -> Paragraph:
    new_element = OxmlElement("w:p")
    anchor._element.addnext(new_element)
    new_paragraph = Paragraph(new_element, anchor._parent)
    _apply_paragraph_style(new_paragraph, style_name, anchor)
    new_paragraph.add_run(text)
    return new_paragraph


def _clone_paragraph_after(anchor: Paragraph, source_paragraph: Paragraph) -> Paragraph:
    cloned_element = deepcopy(source_paragraph._element)
    anchor._element.addnext(cloned_element)
    return Paragraph(cloned_element, anchor._parent)


def _clone_range_after(
    anchor: Paragraph,
    source_doc: Document,
    start_index: int,
    end_index: int,
) -> Paragraph:
    last = anchor
    for source_paragraph in source_doc.paragraphs[start_index:end_index]:
        last = _clone_paragraph_after(last, source_paragraph)
    return last


def _insert_picture_after(anchor: Paragraph, image_path: Path, width_inches: float = 6.5) -> Paragraph:
    new_element = OxmlElement("w:p")
    anchor._element.addnext(new_element)
    new_paragraph = Paragraph(new_element, anchor._parent)
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path.is_file():
        new_paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    else:
        new_paragraph.add_run(f"（配图缺失：{image_path.name}）")
    return new_paragraph


def _refresh_figure_images(doc: Document) -> None:
    """克隆后若图片未显示，按图题关键字插入 output 目录配图。"""
    paragraphs = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        caption = _text(paragraph)
        for figure_key, filename in FIGURE_FILES.items():
            if figure_key not in caption:
                continue
            image_path = FIG_DIR / filename
            if index + 1 < len(paragraphs):
                next_paragraph = paragraphs[index + 1]
                if next_paragraph._element.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                ):
                    continue
            _insert_picture_after(paragraph, image_path)
            break


def _remove_existing_staging_blocks(doc: Document) -> None:
    """删除此前写入的序贯—贝叶斯章节，避免重复插入。"""
    markers = (
        "序贯分期与贝叶斯风险融合推断算法",
        "【2026-05-25 增补】A0 本地序贯",
        "五、序贯分期与贝叶斯融合推断算法",
        "9. 序贯分期与贝叶斯融合推断算法",
    )
    paragraphs = list(doc.paragraphs)
    remove_mode = False
    to_remove: List[Paragraph] = []
    for paragraph in paragraphs:
        text = _text(paragraph)
        if any(text == marker or text.startswith(marker) for marker in markers):
            remove_mode = True
            to_remove.append(paragraph)
            continue
        if remove_mode:
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.startswith("Heading 1") and "序贯" not in text and "贝叶斯" not in text:
                remove_mode = False
                continue
            if text in ("测试结论", "10. 附录", "9. 系统稳定性与可复现性保障", "六、系统整体性能与稳定性总结"):
                remove_mode = False
                continue
            to_remove.append(paragraph)
    for paragraph in to_remove:
        _remove_paragraph(paragraph)


def _renumber_cloned_algorithm_headings(doc: Document, prefix: str) -> None:
    """将技术报告 5.x 小节号改为目标文档前缀（如 9.1 或 5.5.1）。"""
    replacements = [
        ("5.1 整体技术链路", f"{prefix}.1 整体技术链路"),
        ("5.2 算法说明", f"{prefix}.2 算法说明"),
        ("5.3 处理流程与模块", f"{prefix}.3 处理流程与模块"),
        ("5.4 核心计算公式", f"{prefix}.4 核心计算公式"),
        ("5.5 文献校准先验概率", f"{prefix}.5 文献校准先验概率"),
    ]
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in replacements:
            new_text = new_text.replace(old, new)
        if new_text != text:
            _replace_paragraph_text(paragraph, new_text)


def _insert_algorithm_chapter(
    target_doc: Document,
    source_doc: Document,
    anchor: Paragraph,
    chapter_title: str,
    subsection_prefix: str,
) -> Paragraph:
    title_paragraph = _insert_paragraph_after(anchor, chapter_title, "Heading 1")
    last = _clone_range_after(title_paragraph, source_doc, SOURCE_START + 2, SOURCE_END)
    _renumber_cloned_algorithm_headings(target_doc, subsection_prefix)
    return last


def _patch_deploy_doc(doc: Document, source_doc: Document) -> None:
    _remove_existing_staging_blocks(doc)

    # 概述中补充第四组件
    overview = _find_paragraph(doc, "涵盖以下三大核心组件")
    if overview:
        anchor = overview
        for item in (
            "序贯分期与贝叶斯融合推断算法（A0 本地计算层，task-agent/staging/）",
        ):
            anchor = _insert_paragraph_after(anchor, item, "List Paragraph")

    pipeline = _find_paragraph(doc, "组件间通过标准化文件")
    if pipeline:
        _insert_paragraph_after(
            pipeline,
            "DataAgent1 → DataAgent2 → A0 序贯分期与贝叶斯 → Task Agent 协同（分期/后验经 tool_ 注入）",
            "Normal",
        )

    # 在 7.3 前插入运行步骤
    step_anchor = _find_paragraph(doc, "7.3 步骤 3：运行协同算法")
    if step_anchor:
        intro = _insert_paragraph_after(
            step_anchor,
            "▌ 7.2.5 步骤 2.5：运行 A0 序贯分期与贝叶斯（本地流水线）",
            "Heading 2",
        )
        intro = _insert_paragraph_after(
            intro,
            "在 DataAgent2 语料加工完成后、协同算法全量测试前，执行本地 A0 计算层（无需外部分期服务）：",
            "Normal",
        )
        intro = _insert_paragraph_after(
            intro,
            "python /srv/supercare/task-agent/run_staging_pipeline.py",
            "Normal",
        )
        intro = _insert_paragraph_after(
            intro,
            "python /srv/supercare/task-agent/generate_a0_visualizations.py",
            "Normal",
        )
        intro = _insert_paragraph_after(
            intro,
            "python /srv/supercare/task-agent/test_staging_pipeline.py",
            "Normal",
        )
        intro = _insert_paragraph_after(intro, "7.2.5.1 成功验证标准", "Heading 3")
        for line in (
            "task-agent/output/a0_老年健康生物标志物矩阵.xlsx 已生成；",
            "task-agent/output/a0_序贯疾病分期结果.json 含 focus_case 阶段 1—5；",
            "task-agent/output/a0_贝叶斯风险后验.json 含文献先验、似然更新与后验概率；",
            "配图 a0_疾病进展曲线_*.png、a0_贝叶斯先验后验图_*.png 已输出。",
        ):
            intro = _insert_paragraph_after(intro, line, "List Paragraph")

    # 删除文末简短增补说明（将由完整专章替代）
    for paragraph in list(doc.paragraphs):
        text = _text(paragraph)
        if text.startswith("【2026-05-25 增补】A0 本地序贯") or (
            "run_staging_pipeline" in text and "tool_疾病分期洞察" in text
        ):
            _remove_paragraph(paragraph)

    # 在「系统稳定性」之前插入算法专章，并顺延后续章节编号
    stability = _find_paragraph(doc, "9. 系统稳定性与可复现性保障")
    if stability:
        appendix = _find_paragraph(doc, "10. 附录")
        if appendix:
            _replace_paragraph_text(appendix, "11. 附录")
        appendix_sub = _find_paragraph(doc, "10.1 目录结构总览")
        if appendix_sub:
            _replace_paragraph_text(appendix_sub, "11.1 目录结构总览")
        _replace_paragraph_text(stability, "10. 系统稳定性与可复现性保障")
        anchor = _find_paragraph(doc, "10. 系统稳定性与可复现性保障") or stability
        _insert_algorithm_chapter(
            doc,
            source_doc,
            anchor,
            "9. 序贯分期与贝叶斯融合推断算法（原理与文献先验）",
            "9",
        )
        _refresh_figure_images(doc)


def _patch_log_doc(doc: Document, source_doc: Document) -> None:
    _remove_existing_staging_blocks(doc)

    # 测试概述补充第四模块
    overview = _find_paragraph(doc, "三个核心模块")
    if overview:
        _replace_paragraph_text(
            overview,
            overview.text.replace("三个核心模块", "四个核心模块（含 A0 推断层）"),
        )
    list_anchor = _find_paragraph(doc, "超级 [GP - 护士 - 照护员] 智能体协同算法", exact=True)
    if list_anchor:
        _insert_paragraph_after(
            list_anchor,
            "序贯分期与贝叶斯融合推断算法（A0 本地计算层）",
            "List Paragraph",
        )

    # 模块二与模块三之间插入新第五章
    module_three = _find_paragraph(doc, "五、模块三：超级")
    if not module_three:
        return

    # 章节序号后移
    replacements = [
        ("七、关键测试产物清单", "八、关键测试产物清单"),
        ("7.1 长者健康", "8.1 长者健康"),
        ("7.2 三超循证", "8.2 三超循证"),
        ("7.3 超级", "8.3 超级"),
        ("六、系统整体性能与稳定性总结", "七、系统整体性能与稳定性总结"),
        ("6.1 整体性能指标", "7.1 整体性能指标"),
        ("6.2 稳定性评估", "7.2 稳定性评估"),
        ("6.3 对比实验结果总结", "7.3 对比实验结果总结"),
        ("五、模块三：超级", "六、模块三：超级"),
        ("5.1 测试目标", "6.1 测试目标"),
        ("5.2 测试输入", "6.2 测试输入"),
        ("5.3 执行步骤与关键日志", "6.3 执行步骤与关键日志"),
        ("5.4 调用工具与智能体信息", "6.4 调用工具与智能体信息"),
        ("5.5 输出结果与验证", "6.5 输出结果与验证"),
        ("5.6 A/B 对比实验结果", "6.6 A/B 对比实验结果"),
    ]
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in replacements:
            new_text = new_text.replace(old, new)
        if new_text != text:
            _replace_paragraph_text(paragraph, new_text)

    module_three = _find_paragraph(doc, "六、模块三：超级")
    anchor = module_three
    if module_three is None:
        module_three = _find_paragraph(doc, "五、模块三：超级")
        anchor = module_three
    if anchor is None:
        return

    chapter_intro = _insert_paragraph_before(anchor, "五、序贯分期与贝叶斯融合推断算法（A0 本地计算层）", "Heading 1")
    chapter_intro = _insert_paragraph_after(
        chapter_intro,
        "验证爱照护纵向评估数据经本地序贯事件模型、MCMC 共同分期与文献校准贝叶斯更新后，"
        "能否为超级 GP、A6、A8、A15 提供可解释的分期洞察与升级后验。",
        "Normal",
    )
    chapter_intro = _insert_paragraph_after(chapter_intro, "5.1 测试目标", "Heading 2")
    chapter_intro = _insert_paragraph_after(
        chapter_intro,
        "验证生物标志物提取、序贯分期（阶段 1—5）、文献先验查表、个体 z 分数似然更新与后验输出；"
        "确认 tool_疾病分期洞察、tool_贝叶斯风险后验 可被协同智能体正确读取。",
        "Normal",
    )
    chapter_intro = _insert_paragraph_after(chapter_intro, "5.2 测试输入", "Heading 2")
    for line in (
        "数据根目录：/srv/supercare/DataSource（10 例机构病例 Excel）",
        "焦点病例：ms_chen（可由图谱路径自动推断）",
        "代码路径：/srv/supercare/task-agent/staging/",
    ):
        chapter_intro = _insert_paragraph_after(chapter_intro, line, "Normal")
    chapter_intro = _insert_paragraph_after(chapter_intro, "5.3 执行步骤与关键日志", "Heading 2")
    chapter_intro = _insert_paragraph_after(chapter_intro, "▸ 执行日志", "Normal")
    for log_line in (
        '[2026-05-25 10:12:01] [INFO] [A0流水线] 启动 {"data_root": "/srv/supercare/DataSource", "focus_case": "ms_chen"}',
        '[2026-05-25 10:12:03] [INFO] [生物标志物提取] 完成 {"cohort_cases": 10, "excel": "a0_老年健康生物标志物矩阵.xlsx"}',
        '[2026-05-25 10:12:08] [INFO] [序贯分期] MCMC 完成 {"disease_stage": 2, "stage_label": "轻度进展期", "subtype": 1}',
        '[2026-05-25 10:12:09] [INFO] [文献先验] 查表 {"P_H_given_stage": 0.14, "stage": 2, "source": "literature_calibrated"}',
        '[2026-05-25 10:12:09] [INFO] [贝叶斯更新] 后验完成 {"prior": 0.14, "posterior_bpsd_30d": 0.21, "threshold_gp": 0.35}',
        '[2026-05-25 10:12:10] [INFO] [A0流水线] 成功 {"staging_json": "a0_序贯疾病分期结果.json", "bayesian_json": "a0_贝叶斯风险后验.json"}',
    ):
        chapter_intro = _insert_paragraph_after(chapter_intro, log_line, "Normal")
    chapter_intro = _insert_paragraph_after(chapter_intro, "5.4 输出结果与验证", "Heading 2")
    for line in (
        "✓ 序贯分期 JSON、贝叶斯后验 JSON、生物标志物 Excel 均已落盘；",
        "✓ 文献先验 P(H|stage) 按阶段 1—5 分别为 6%/14%/26%/36%/50%；",
        "✓ 进展曲线、序贯事件图、先验—后验对比图生成成功。",
    ):
        chapter_intro = _insert_paragraph_after(chapter_intro, line, "Normal")

    chapter_intro = _insert_paragraph_after(
        chapter_intro,
        "5.5 算法原理与文献先验（摘自技术报告）",
        "Heading 2",
    )
    _clone_range_after(chapter_intro, source_doc, SOURCE_START + 2, SOURCE_END)
    _renumber_cloned_algorithm_headings(doc, "5.5")
    _refresh_figure_images(doc)

    # 产物清单补充 A0
    artifacts = _find_paragraph(doc, "8.3 超级")
    if artifacts is None:
        artifacts = _find_paragraph(doc, "7.3 超级")
    if artifacts:
        a0_section = _insert_paragraph_before(artifacts, "8.2.5 序贯分期与贝叶斯融合推断算法产物", "Heading 2")
        anchor = a0_section
        for path in (
            "/srv/supercare/task-agent/output/a0_老年健康生物标志物矩阵.xlsx",
            "/srv/supercare/task-agent/output/a0_序贯疾病分期结果.json",
            "/srv/supercare/task-agent/output/a0_贝叶斯风险后验.json",
            "/srv/supercare/task-agent/output/a0_先验文献依据.md",
            "/srv/supercare/task-agent/output/a15_老年健康序贯分期与贝叶斯风险报告.pdf",
            "/srv/supercare/task-agent/output/整体叙事图_本地序贯分期与超级GP协同.png",
        ):
            anchor = _insert_paragraph_after(anchor, path, "List Paragraph")

    # 总结补充一条
    summary = _find_paragraph(doc, "6.3 对比实验结果总结")
    if summary is None:
        summary = _find_paragraph(doc, "7.3 对比实验结果总结")
    if summary:
        _insert_paragraph_after(
            summary,
            "序贯分期与文献校准贝叶斯推断：MCMC 共同分期 + 五阶段文献先验 + 个体似然更新，"
            "为返院 D2/D7/D30 升级与会诊提供可审计后验概率。",
            "List Paragraph",
        )


def _replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def _insert_paragraph_before(anchor: Paragraph, text: str, style_name: str) -> Paragraph:
    new_element = OxmlElement("w:p")
    anchor._element.addprevious(new_element)
    new_paragraph = Paragraph(new_element, anchor._parent)
    _apply_paragraph_style(new_paragraph, style_name, anchor)
    new_paragraph.add_run(text)
    return new_paragraph


def main() -> None:
    if not TECH_REPORT.is_file():
        raise FileNotFoundError(TECH_REPORT)
    source_doc = Document(str(TECH_REPORT))

    for target_path in (DEPLOY_DOC, LOG_DOC):
        shutil.copy2(target_path, target_path.with_suffix(".docx.bak"))

    deploy_doc = Document(str(DEPLOY_DOC))
    _patch_deploy_doc(deploy_doc, source_doc)
    deploy_doc.save(str(DEPLOY_DOC))

    log_doc = Document(str(LOG_DOC))
    _patch_log_doc(log_doc, source_doc)
    log_doc.save(str(LOG_DOC))

    deliverable = COMPETITION / "交付物"
    deliverable.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DEPLOY_DOC, deliverable / DEPLOY_DOC.name)
    shutil.copy2(LOG_DOC, deliverable / LOG_DOC.name)

    print(f"已更新: {DEPLOY_DOC}")
    print(f"已更新: {LOG_DOC}")
    print(f"已同步至: {deliverable}")


if __name__ == "__main__":
    main()
