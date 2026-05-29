#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将序贯分期与贝叶斯内容重组为「五、面向 BPSD 返院照护的序贯分期与贝叶斯风险融合算法」：
公式左式右号、变量说明、图注解释、参考文献移至文末。
"""

from __future__ import annotations

import shutil
from pathlib import Path
from typing import List, Optional, Tuple, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from docx_equations import (
    EQUATION_LOGISTIC_LEVEL,
    EQUATION_LOG_LIKELIHOOD,
    EQUATION_LR_Z_NEG,
    EQUATION_LR_Z_POS,
    EQUATION_ODDS,
    EQUATION_ODDS_UPDATE,
    EQUATION_POSTERIOR,
    EQUATION_PRIOR_STAGE,
    EQUATION_ZSCORE,
    MATH_NS,
    mathml_to_omml,
)
from staging.prior_literature import (
    PRIOR_ASSUMPTION_H,
    PRIOR_FORMULA_NOTE,
    PRIOR_REFERENCES,
    STAGE_BPSD_PRIORS_LITERATURE,
    STAGE_LABELS,
    STAGE_PRIOR_DERIVATION,
)

COMPETITION = Path("/srv/supercare/比赛文档")
REPORT_PATH = (
    COMPETITION / "1_技术报告_超级[GP-护士-照护员]智能体协同算法_含序贯贝叶斯_20260525.docx"
)
DELIVERABLE = COMPETITION / "交付物" / REPORT_PATH.name
TASK_OUTPUT = Path(__file__).resolve().parent / "output"

FIGURES = [
    (
        "图5-1",
        "整体技术链路：爱照护纵向评估 → 序贯分期推断 → 贝叶斯风险融合 → 超级GP协同决策",
        "图5-1_整体技术链路_爱照护序贯贝叶斯超级GP.png",
        "图5-1 从数据到决策的整体技术链路。输入为爱照护专业照护机构长期积累的 NPI、认知与 ADL 等纵向评估；"
        "中部经序贯分期估计当前 1—5 级疾病阶段，再经文献校准贝叶斯层得到 BPSD 升级后验；"
        "右侧将分期洞察与风险后验注入超级 GP、护士与照护员协同链，在返院关键节点触发升级与会诊。",
    ),
    (
        "图5-2",
        "疾病进展曲线（横轴为共同分期，多标志物同图）",
        "a0_疾病进展曲线_陈女士.png",
        "图5-2 疾病进展曲线。横轴为共同疾病分期（1—5 级），纵轴为各标志物归一化水平，"
        "展示不同指标随分期推进由基线向异常过渡的趋势。",
    ),
    (
        "图5-3",
        "序贯事件与 MCMC 整体分期",
        "a0_序贯事件进展图_陈女士.png",
        "图5-3 序贯事件与 MCMC 分期示意。展示各标志物异常触发顺序及共同分期估计；"
        "MCMC 在爱照护队列最新截面上给出当前阶段，作为贝叶斯先验查表依据。",
    ),
    (
        "图5-4",
        "贝叶斯先验与后验",
        "a0_贝叶斯先验后验图_陈女士.png",
        "图5-4 贝叶斯先验与后验对比。左柱为按 MCMC 分期查文献得到的先验，右柱为融合个体 z 分数后的后验；"
        "用于返院 D2/D7/D30 会诊与升级阈值判断。",
    ),
]

EQUATION_SPECS = [
    {
        "mathml": EQUATION_LOGISTIC_LEVEL,
        "label": "（公式1）",
        "vars": (
            "式中：k 为标志物编号；L_k 为标志物 k 的水平；s 为共同疾病分期（1—5）；"
            "b_k 为正常基线；a_k 为异常后水平；τ_k 为异常出现的分期时刻；w_k 为 logistic 过渡宽度。"
        ),
    },
    {
        "mathml": EQUATION_LOG_LIKELIHOOD,
        "label": "（公式2）",
        "vars": (
            "式中：D 为当前评估观测；s 为分期；x_k 为实测值；E_k 为公式（1）的期望；σ_k 为残差尺度。"
        ),
    },
    {
        "mathml": EQUATION_PRIOR_STAGE,
        "label": "（公式3）",
        "vars": (
            "式中：H 为未来 30 天内 BPSD 临床显著恶化；p_s、λ_s 分别为该期患病率与月尺度加重比例[1-3]。"
        ),
    },
    {
        "mathml": EQUATION_ZSCORE,
        "label": "（公式4）",
        "vars": "式中：z_i 为标准化偏离；x_i 为个体值；μ_queue、σ_queue 为爱照护队列均值与标准差。",
    },
    {
        "mathml": EQUATION_LR_Z_POS,
        "label": "（公式5a）",
        "vars": "式中：LR(z_i) 为似然比；LR_0 为基准；z_cap 为截断上限；适用于 z_i>0。",
    },
    {
        "mathml": EQUATION_LR_Z_NEG,
        "label": "（公式5b）",
        "vars": "式中：符号同公式（5a）；适用于 z_i<0。",
    },
    {
        "mathml": EQUATION_ODDS,
        "label": "（公式6）",
        "vars": "式中：odds(H) 为先验优势比；P(H) 由公式（3）或文献查表得到。",
    },
    {
        "mathml": EQUATION_ODDS_UPDATE,
        "label": "（公式7）",
        "vars": "式中：D 为个体观测；n 为标志物个数；LR(z_i) 为各指标似然比。",
    },
    {
        "mathml": EQUATION_POSTERIOR,
        "label": "（公式8）",
        "vars": "式中：P(H|D) 为后验恶化概率，用于 GP 会诊与事件升级。",
    },
]

STAGE_CITES = {1: "[1]", 2: "[1][5]", 3: "[3]", 4: "[2][4]", 5: "[2][6]"}
STAGE_REF_NOTES = {1: "1", 2: "1、5", 3: "3", 4: "2、4", 5: "2、6"}


def _text(paragraph: Paragraph) -> str:
    return (paragraph.text or "").strip()


def _find(doc: Document, keyword: str, *, exact: bool = False) -> Optional[Paragraph]:
    for p in doc.paragraphs:
        t = _text(p)
        if exact:
            if t != keyword:
                continue
        elif keyword not in t:
            continue
        return p
    return None


def _remove(p: Paragraph) -> None:
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _insert_before(anchor: Paragraph, text: str = "", style: Optional[str] = None) -> Paragraph:
    new_el = OxmlElement("w:p")
    anchor._element.addprevious(new_el)
    np = Paragraph(new_el, anchor._parent)
    if style:
        np.style = style
    if text:
        np.add_run(text)
    return np


def _insert_blocks_before(anchor: Paragraph, blocks: List[Tuple[str, str]]) -> None:
    for style, text in reversed(blocks):
        anchor = _insert_before(anchor, text, style)


def _append_omml(paragraph: Paragraph, mathml: str) -> None:
    inner = mathml_to_omml(mathml.strip())
    if not inner.startswith("<m:oMath"):
        inner = f"<m:oMath>{inner}</m:oMath>"
    paragraph._p.append(parse_xml(f'<m:oMathPara xmlns:m="{MATH_NS}">{inner}</m:oMathPara>'))


def _make_equation_table(document: Document, mathml: str, label: str):
    table = document.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells[0], table.rows[0].cells[1]
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _append_omml(lp, mathml)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = rp.add_run(label)
    run.font.size = Pt(10.5)
    return table


def _insert_table_after(paragraph: Paragraph, table) -> None:
    paragraph._element.addnext(table._tbl)


def _insert_picture_before(anchor: Paragraph, caption: str, image: Path) -> None:
    cap = _insert_before(anchor, caption, "Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.bold = True
    if image.is_file():
        pic = _insert_before(anchor, "")
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(image), width=Inches(5.6))
    else:
        _insert_before(anchor, f"（配图缺失：{image.name}）", "Normal")


def _remove_old_a0_block(doc: Document) -> None:
    """删除文末或交付成果前遗留的图 A0 专章块。"""
    markers = ("图A0-4", "图A0-3", "图A0-2", "图A0-1", "本节阐述面向 BPSD 返院照护的序贯分期与贝叶斯风险融合算法原理")
    for _ in range(3):
        start = None
        for m in markers:
            start = _find(doc, m)
            if start:
                break
        if not start:
            break
        end = _find(doc, "交付成果与行业价值") or _find(doc, "参考文献", exact=True)
        if not end:
            break
        removing, to_del = False, []
        for p in doc.paragraphs:
            if p == start:
                removing = True
            if p == end:
                break
            if removing:
                to_del.append(p)
        if not to_del:
            _remove(start)
            continue
        for p in to_del:
            _remove(p)


def _remove_mid_references(doc: Document) -> None:
    h = _find(doc, "参考文献（先验概率）", exact=True)
    if not h:
        return
    to_del, on = [], False
    for p in doc.paragraphs:
        if p == h:
            on = True
        if on:
            sn = p.style.name if p.style else ""
            if on and p != h and sn == "Heading 1":
                break
            to_del.append(p)
    for p in to_del:
        _remove(p)


def _build_section_blocks() -> List[Union[Tuple[str, str], Tuple[str, dict]]]:
    blocks: List[Union[Tuple[str, str], Tuple[str, dict]]] = [
        ("Heading 2", "五、面向 BPSD 返院照护的序贯分期与贝叶斯风险融合算法"),
        (
            "Body Text",
            "本节阐述序贯分期与贝叶斯风险融合推断算法的原理、核心公式、文献先验来源与示意配图；"
            "配图为爱照护流水线运行示例，不代表单一病例临床报告。",
        ),
        ("Heading 3", "5.1 整体技术链路"),
        (
            "Body Text",
            "整体链路分为四层：① 爱照护专业照护机构纵向评估数据（NPI、MMSE/MoCA、Barthel 等）；"
            "② 序贯事件建模与 MCMC 共同分期；③ 文献校准贝叶斯后验；"
            "④ 超级 GP—护士—照护员协同决策。量化结果经「疾病分期洞察」「风险后验」工具"
            "注入返院 D2/D7/D30 等节点。",
        ),
        ("fig", {"idx": 0}),
        ("Heading 3", "5.2 算法说明"),
        (
            "Body Text",
            "「序贯」含义：各 BPSD 相关指标随病程逐步由正常转向异常，并在同一套 1—5 级分期标尺上"
            "刻画当前所处阶段，避免对各指标分别分期后再简单平均。",
        ),
        (
            "Body Text",
            "每个标志物 k 用序贯事件模型描述：正常基线 b_k、异常水平 a_k、异常出现时刻 τ_k；"
            "分期越靠后，各指标异常出现的时间通常越晚。",
        ),
        (
            "Body Text",
            "Metropolis-Hastings MCMC 在爱照护队列最新评估截面上推断当前疾病分期；"
            "疾病进展曲线横轴为共同分期，纵轴为标志物水平（多标志物同图时采用归一化进展指数）。",
        ),
        (
            "Body Text",
            "贝叶斯层关注事件 H：未来 30 天内 BPSD 是否临床显著恶化（NPI 总分升高≥4 或任一分项≥4）[1]。"
            "先验 P(H) 由 MCMC 分期对照文献校准表确定[1-3]；各标志物相对爱照护队列的 z 分数"
            "转换为似然比后更新 odds，得到后验概率，并与 GP 会诊阈值（35%/25%/20%）联动。",
        ),
        (
            "Body Text",
            "推断流水线在爱照护机构侧部署运行；算法细节另见《序贯分期与贝叶斯风险融合推断算法说明》专册。",
        ),
        ("Heading 3", "5.3 处理流程与模块"),
        (
            "Body Text",
            "处理链路依次为：生物标志物提取 → 序贯事件与 MCMC 分期 → 文献先验构建 → 贝叶斯后验更新。",
        ),
        (
            "Body Text",
            "输出包括序贯分期结构化结果、风险后验结果及疾病进展与先验—后验对比配图，供协同智能体调用。",
        ),
        ("Heading 3", "5.4 核心计算公式"),
        ("Body Text", "下列公式采用左式右号排版；各式后附主要变量说明。"),
    ]
    for spec in EQUATION_SPECS:
        blocks.append(("equation", spec))
    blocks.extend(
        [
            ("Body Text", PRIOR_ASSUMPTION_H),
            ("Body Text", PRIOR_FORMULA_NOTE + " 详见文末参考文献[1-6]。"),
            ("Heading 3", "5.5 文献校准先验概率"),
        ]
    )
    for stage in range(1, 6):
        blocks.append(
            (
                "Body Text",
                f"阶段{stage}（{STAGE_LABELS[stage]}）：P(H|stage)="
                f"{STAGE_BPSD_PRIORS_LITERATURE[stage]:.0%}{STAGE_CITES[stage]} — "
                f"{STAGE_PRIOR_DERIVATION[stage]}（文献{STAGE_REF_NOTES[stage]}）",
            )
        )
    for idx in range(1, 4):
        blocks.append(("fig", {"idx": idx}))
    return blocks


def _remove_duplicate_references(doc: Document) -> None:
    """保留文末参考文献，删除交付成果前的重复条目。"""
    end_refs = _find(doc, "参考文献", exact=True)
    if not end_refs:
        return
    for p in list(doc.paragraphs):
        t = _text(p)
        if p == end_refs:
            break
        if t.startswith("[1] Lyketsos") or t.startswith("[2] Zhao"):
            _remove(p)


def _rebuild_equation_block(doc: Document) -> None:
    """5.4 节：清空后重建为「公式表 + 式中说明」交替排列。"""
    intro = _find(doc, "下列公式采用左式右号排版")
    prior_h = _find(doc, "5.5 文献校准先验概率", exact=True)
    if not intro or not prior_h:
        return

    intro_el, prior_el = intro._element, prior_h._element
    el = intro_el.getnext()
    while el is not None and el != prior_el:
        nxt = el.getnext()
        intro_el.getparent().remove(el)
        el = nxt

    anchor_el = intro_el
    for spec in EQUATION_SPECS:
        tbl = _make_equation_table(doc, spec["mathml"], spec["label"])
        anchor_el.addnext(tbl._tbl)
        anchor_el = tbl._tbl
        p_el = OxmlElement("w:p")
        anchor_el.addnext(p_el)
        vp = Paragraph(p_el, intro._parent)
        vp.style = "Body Text"
        vp.add_run(spec["vars"])
        anchor_el = p_el

    for extra in (PRIOR_ASSUMPTION_H, PRIOR_FORMULA_NOTE + " 详见文末参考文献[1-6]。"):
        p_el = OxmlElement("w:p")
        anchor_el.addnext(p_el)
        bp = Paragraph(p_el, intro._parent)
        bp.style = "Body Text"
        bp.add_run(extra)
        anchor_el = p_el


def _dedupe_reference_section(doc: Document) -> None:
    """仅保留最后一处「参考文献」章节，并删除正文中的孤立条目。"""
    headings = [p for p in doc.paragraphs if _text(p) == "参考文献"]
    if len(headings) > 1:
        for heading in headings[:-1]:
            el = heading._element.getnext()
            while el is not None:
                nxt = el.getnext()
                tag = el.tag.split("}")[-1]
                if tag == "p":
                    para = Paragraph(el, heading._parent)
                    sn = para.style.name if para.style else ""
                    if sn == "Heading 1":
                        break
                    t = _text(para)
                    if t.startswith("[") or not t:
                        el.getparent().remove(el)
                el = nxt
            _remove(heading)

    last_ref = _find(doc, "参考文献", exact=True)
    if not last_ref:
        return
    for p in list(doc.paragraphs):
        if p == last_ref:
            break
        t = _text(p)
        if t.startswith("[1] Lyketsos") or t.startswith("[2] Zhao"):
            _remove(p)


def _insert_section5(doc: Document, anchor: Paragraph) -> None:
    """在「典型任务执行示例」前插入第五节（倒序插入文本块，再正向挂接公式表与图）。"""
    blocks = _build_section_blocks()
    text_blocks: List[Tuple[str, str]] = []
    equations: List[dict] = []
    figures: List[int] = []

    for item in blocks:
        if item[0] == "equation":
            equations.append(item[1])  # type: ignore
        elif item[0] == "fig":
            figures.append(item[1]["idx"])  # type: ignore
        else:
            text_blocks.append(item)  # type: ignore

    _insert_blocks_before(anchor, text_blocks)

    # 附图：插在对应 fig 占位 — 按 fig idx 在 5.2 前、5.5 后
    h52 = _find(doc, "5.2 算法说明", exact=True)
    if h52 and 0 in figures:
        f = FIGURES[0]
        _insert_picture_before(h52, f"{f[0]} {f[1]}", TASK_OUTPUT / f[2])
        _insert_before(h52, f[3], "Body Text")

    h55 = _find(doc, "5.5 文献校准先验概率", exact=True)
    if h55:
        insert_point = h55
        for idx in sorted([i for i in figures if i > 0], reverse=True):
            f = FIGURES[idx]
            # 在 5.5 标题之后插入：找阶段5段落后
            insert_point = h55
            for p in doc.paragraphs:
                if _text(p).startswith("阶段5（"):
                    insert_point = p
                    break
            nxt = insert_point._element.getnext()
            anchor_f = Paragraph(nxt, insert_point._parent) if nxt is not None else insert_point
            _insert_picture_before(anchor_f, f"{f[0]} {f[1]}", TASK_OUTPUT / f[2])
            _insert_before(anchor_f, f[3], "Body Text")


def _append_references(doc: Document) -> None:
    if _find(doc, "参考文献", exact=True) and _find(doc, "[1] Lyketsos"):
        return
    concl = _find(doc, "结论与展望") or doc.paragraphs[-1]
    blocks = [("Heading 1", "参考文献")]
    for ref in PRIOR_REFERENCES:
        blocks.append(("Body Text", f"[{ref['id']}] {ref['citation']}"))
    _insert_blocks_before(concl, blocks)


def _patch_crossref(doc: Document) -> None:
    for p in doc.paragraphs:
        t = _text(p)
        if "见本章后文「序贯分期与贝叶斯风险融合」专节" in t:
            new = t.replace(
                "见本章后文「序贯分期与贝叶斯风险融合」专节",
                "详见「五、面向 BPSD 返院照护的序贯分期与贝叶斯风险融合算法」",
            )
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""


def reorganize(path: Path) -> None:
    doc = Document(str(path))
    _remove_old_a0_block(doc)
    _remove_mid_references(doc)
    _remove_duplicate_references(doc)

    anchor = _find(doc, "典型任务执行示例") or _find(doc, "交付成果与行业价值")
    if not anchor:
        raise RuntimeError("未找到第五节插入锚点")

    if not _find(doc, "五、面向 BPSD 返院照护的序贯分期与贝叶斯风险融合算法", exact=True):
        _insert_section5(doc, anchor)

    _rebuild_equation_block(doc)
    _remove_old_a0_block(doc)
    _remove_duplicate_references(doc)
    _dedupe_reference_section(doc)
    _append_references(doc)
    _dedupe_reference_section(doc)
    _patch_crossref(doc)
    doc.save(str(path))


def main() -> None:
    reorganize(REPORT_PATH)
    DELIVERABLE.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REPORT_PATH, DELIVERABLE)
    print(f"已重组: {REPORT_PATH}")


if __name__ == "__main__":
    main()
