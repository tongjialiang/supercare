#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从原版技术报告 docx 另存为「含序贯贝叶斯」版：在摘要、创新、技术方案等章节增补 A0 内容，
并新增专章（算法、OMML 公式、配图、文献先验）。不覆盖原版 docx。
"""

from __future__ import annotations

import shutil
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from docx_equations import (
    EQUATION_LOGISTIC_LEVEL,
    EQUATION_LOG_LIKELIHOOD,
    EQUATION_LR_Z_NEG,
    EQUATION_LR_Z_POS,
    EQUATION_ODDS,
    EQUATION_ODDS_UPDATE,
    EQUATION_POSTERIOR,
    EQUATION_PRIOR_STAGE,
    EQUATION_ZSCORE,
    MATH_NS,
    mathml_to_omml,
)
from staging.prior_literature import (
    PRIOR_ASSUMPTION_H,
    PRIOR_FORMULA_NOTE,
    PRIOR_REFERENCES,
    STAGE_BPSD_PRIORS_LITERATURE,
    STAGE_LABELS,
    STAGE_PRIOR_DERIVATION,
)

COMPETITION = Path("/srv/supercare/比赛文档")
SOURCE_REPORT = COMPETITION / "1_技术报告_超级[GP-护士-照护员]智能体协同算法.docx"
OUTPUT_REPORT = COMPETITION / "1_技术报告_超级[GP-护士-照护员]智能体协同算法_含序贯贝叶斯_20260525.docx"
TASK_OUTPUT = Path(__file__).resolve().parent / "output"
DELIVERABLE = COMPETITION / "交付物" / OUTPUT_REPORT.name

FIGURES = [
    ("图A0-1 整体叙事：DataSource → 序贯分期 → 贝叶斯 → 超级GP", "整体叙事图_本地序贯分期与超级GP协同.png"),
    ("图A0-2 疾病进展曲线（横轴=共同分期，多标志物同图）", "a0_疾病进展曲线_陈女士.png"),
    ("图A0-3 序贯事件与 MCMC 整体分期", "a0_序贯事件进展图_陈女士.png"),
    ("图A0-4 贝叶斯先验与后验", "a0_贝叶斯先验后验图_陈女士.png"),
]


def _insert_paragraph_before(
    anchor: Paragraph,
    text: str = "",
    style: Optional[str] = None,
) -> Paragraph:
    """在 anchor 段落之前插入新段落。"""
    new_element = OxmlElement("w:p")
    anchor._element.addprevious(new_element)
    new_paragraph = Paragraph(new_element, anchor._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _insert_blocks_before(anchor: Paragraph, blocks: List[Tuple[str, str]]) -> Paragraph:
    """按顺序在 anchor 前插入多段（blocks 顺序保持阅读顺序）。"""
    first_inserted = anchor
    for style_name, text in reversed(blocks):
        first_inserted = _insert_paragraph_before(first_inserted, text, style_name)
    return first_inserted


def _insert_equation_before(anchor: Paragraph, mathml: str, caption: str = "") -> Paragraph:
    """在 anchor 前插入 OMML 公式段。"""
    omml_inner = mathml_to_omml(mathml.strip())
    if omml_inner.startswith("<m:oMath"):
        omath_block = omml_inner
    else:
        omath_block = f"<m:oMath>{omml_inner}</m:oMath>"

    equation_paragraph = _insert_paragraph_before(anchor, "")
    equation_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omath_para_xml = (
        f'<m:oMathPara xmlns:m="{MATH_NS}">'
        f"{omath_block}"
        "</m:oMathPara>"
    )
    equation_paragraph._p.append(parse_xml(omath_para_xml))

    if caption:
        caption_paragraph = _insert_paragraph_before(anchor, caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption_paragraph.runs:
            run.italic = True
    return equation_paragraph


def _insert_picture_before(anchor: Paragraph, caption: str, image_path: Path) -> None:
    caption_paragraph = _insert_paragraph_before(anchor, caption, "Normal")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption_paragraph.runs:
        run.bold = True

    if image_path.is_file():
        picture_paragraph = _insert_paragraph_before(anchor, "")
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))
    else:
        _insert_paragraph_before(anchor, f"（配图缺失：{image_path.name}）", "Normal")


def _find_paragraph(doc: Document, keyword: str, style_contains: str = "") -> Optional[Paragraph]:
    for paragraph in doc.paragraphs:
        text = paragraph.text or ""
        if keyword not in text:
            continue
        if style_contains and style_contains not in (paragraph.style.name or ""):
            continue
        return paragraph
    return None


def _patch_summary_and_innovation(doc: Document) -> None:
    background_heading = _find_paragraph(doc, "项目背景", "Heading")
    if background_heading:
        _insert_blocks_before(
            background_heading,
            [
                (
                    "Body Text",
                    "【A0 本地计算层】在 DataSource 机构队列上提取 BPSD 相关生物标志物（NPI、MMSE/MoCA、Barthel），"
                    "经序贯事件异常模型与 MCMC 推断共同疾病分期，再以文献校准贝叶斯先验融合标志物 z-score 得到 BPSD 升级后验，"
                    "通过 tool_疾病分期洞察、tool_贝叶斯风险后验 注入超级 GP/护士/照护员智能体；全程本地运行。",
                ),
            ],
        )

    data_agent_heading = _find_paragraph(doc, "长者健康计算图DataAgent", "Heading 1")
    if data_agent_heading:
        _insert_blocks_before(
            data_agent_heading,
            [
                ("Heading 3", "A0 本地序贯分期与贝叶斯风险融合"),
                (
                    "Body Text",
                    "创新性提出「序贯共同分期」范式：每个生物标志物对应一个序贯事件（logistic onset），"
                    "全体标志物共用分期轴 1–5，由 MCMC 推断亚型与阶段；贝叶斯层采用国内外文献分层先验 "
                    "P(H|stage)≈p_s×λ_s，再经标志物似然比更新后验，驱动 GP 会诊阈值（35%/25%/20%）。",
                ),
                (
                    "Body Text",
                    "输出：a0_序贯疾病分期结果.json、a0_贝叶斯风险后验.json、疾病进展曲线与先验后验配图；"
                    "详见本章「A0 本地序贯分期与贝叶斯风险融合」。",
                ),
            ],
        )

    gp_bpsd_heading = _find_paragraph(doc, "BPSD患者全周期照护工作流", "Heading")
    if gp_bpsd_heading:
        _insert_blocks_before(
            gp_bpsd_heading,
            [
                ("Heading 2", "A0 本地序贯分期与贝叶斯工具"),
                (
                    "Body Text",
                    "在 11 类标准化图谱工具基础上，新增 tool_疾病分期洞察（MCMC 分期、亚型、序贯事件、进展曲线）"
                    "与 tool_贝叶斯风险后验（文献先验、似然更新、会诊阈值）。返院 D2/D7/D30 等阶段可读取后验概率，"
                    "触发 A8 升级与 A3 会诊建议。",
                ),
            ],
        )

    goals_paragraph = _find_paragraph(doc, "超级[GP—护士—照护员]智能体协同算法：面向业务协同")
    if goals_paragraph:
        goals_paragraph._element.addnext(OxmlElement("w:p"))
        new_paragraph = Paragraph(goals_paragraph._element.getnext(), goals_paragraph._parent)
        new_paragraph.style = "Body Text"
        new_paragraph.add_run(
            "A0 本地序贯分期与贝叶斯层：DataSource → 生物标志物矩阵 → MCMC 分期 → 文献先验贝叶斯后验 → 工具注入协同智能体。"
        )


def _insert_a0_chapter_before_delivery(doc: Document) -> None:
    delivery_heading = _find_paragraph(doc, "交付成果与行业价值", "Heading 1")
    if not delivery_heading:
        delivery_heading = doc.paragraphs[-1]

    anchor = delivery_heading

    def prepend_text(style: str, text: str) -> None:
        nonlocal anchor
        anchor = _insert_paragraph_before(anchor, text, style)

    def prepend_equation(mathml: str, caption: str) -> None:
        nonlocal anchor
        anchor = _insert_equation_before(anchor, mathml, caption)

    def prepend_picture(caption: str, filename: str) -> None:
        nonlocal anchor
        _insert_picture_before(anchor, caption, TASK_OUTPUT / filename)

    # 参考文献（倒序构建）
    for reference in reversed(PRIOR_REFERENCES):
        prepend_text("Body Text", f"[{reference['id']}] {reference['citation']}")
        prepend_text("Body Text", f"    用途：{reference['use']}")
    prepend_text("Heading 2", "参考文献（先验概率）")

    # 配图
    for caption, filename in reversed(FIGURES):
        prepend_picture(caption, filename)

    # 先验表
    for stage_index in reversed(range(1, 6)):
        prepend_text(
            "Body Text",
            f"阶段{stage_index}（{STAGE_LABELS[stage_index]}）："
            f"P(H|stage)={STAGE_BPSD_PRIORS_LITERATURE[stage_index]:.0%} — "
            f"{STAGE_PRIOR_DERIVATION[stage_index]}",
        )
    prepend_text("Heading 2", "文献校准先验概率")
    prepend_text("Body Text", PRIOR_FORMULA_NOTE)
    prepend_text("Body Text", PRIOR_ASSUMPTION_H)

    # 公式
    prepend_equation(EQUATION_POSTERIOR, "式(8) 后验概率")
    prepend_equation(EQUATION_ODDS_UPDATE, "式(7) odds 更新")
    prepend_equation(EQUATION_ODDS, "式(6) 先验 odds")
    prepend_equation(EQUATION_LR_Z_NEG, "式(5b) z_i < 0")
    prepend_equation(EQUATION_LR_Z_POS, "式(5a) z_i > 0")
    prepend_equation(EQUATION_ZSCORE, "式(4) 队列 z-score")
    prepend_equation(EQUATION_PRIOR_STAGE, "式(3) 文献先验")
    prepend_equation(EQUATION_LOG_LIKELIHOOD, "式(2) MCMC 似然")
    prepend_equation(EQUATION_LOGISTIC_LEVEL, "式(1) 序贯事件 logistic 期望水平")
    prepend_text("Heading 2", "核心计算公式")

    # 算法正文
    algorithm_paragraphs = [
        ("Heading 2", "算法说明"),
        (
            "Body Text",
            "「序贯」含义：多个 BPSD 生物标志物按预设顺序依次出现异常，全体标志物处在同一条疾病分期轴 "
            "s∈{1,…,5} 上，而非各指标单独分期再平均。",
        ),
        (
            "Body Text",
            "每个标志物 k 对应序贯事件模型：baseline b_k、异常水平 a_k、onset τ_k；"
            "分期越靠后的事件 onset 越晚。亚型由事件顺序定义（如认知→MoCA→NPI→ADL 等三种轨迹）。",
        ),
        (
            "Body Text",
            "Metropolis-Hastings MCMC 在队列最新截面上联合推断亚型与分期；"
            "疾病进展曲线：横轴为共同分期，纵轴为标志物水平（多标志物同图时采用归一化进展指数）。",
        ),
        (
            "Body Text",
            "贝叶斯层假设 H：30 天内 BPSD 临床显著恶化（NPI 总分升高≥4 或任一分项≥4）。"
            "先验 P(H) 由 MCMC 分期查文献校准表；各标志物队列 z-score 映射为连续似然比 LR，"
            "odds 更新后得到后验，与 GP 阈值联动。",
        ),
        (
            "Body Text",
            "本地命令：python task-agent/run_staging_pipeline.py；"
            "python task-agent/generate_a0_visualizations.py。"
            "算法详述另见：本地序贯分期与贝叶斯风险融合_算法说明.docx。",
        ),
        ("Heading 2", "数据流与模块"),
        ("Body Text", "staging/biomarker_extraction.py → sustain_mcmc_model.py → prior_literature.py → bayesian_risk.py"),
        ("Body Text", "产出：a0_序贯疾病分期结果.json、a0_贝叶斯风险后验.json、进展曲线 PNG。"),
    ]
    for style_name, text in reversed(algorithm_paragraphs):
        prepend_text(style_name, text)

    prepend_text(
        "Body Text",
        "本章说明 A0 本地序贯疾病分期与贝叶斯风险融合的算法原理、公式、文献先验来源与配图；"
        "配图为流水线运行示例，非个体病例临床报告。",
    )
    _insert_paragraph_before(anchor, "A0 本地序贯分期与贝叶斯风险融合", "Heading 1")


def _patch_delivery_section(doc: Document) -> None:
    deliverable_paragraph = _find_paragraph(doc, "技术报告：本技术报告")
    if deliverable_paragraph:
        deliverable_paragraph.add_run(
            " 另含《含序贯贝叶斯》修订版及 A0 JSON/配图、算法说明 docx。"
        )


def main() -> None:
    if not SOURCE_REPORT.is_file():
        raise FileNotFoundError(f"未找到原版技术报告: {SOURCE_REPORT}")

    shutil.copy2(SOURCE_REPORT, OUTPUT_REPORT)
    document = Document(str(OUTPUT_REPORT))

    _patch_summary_and_innovation(document)
    _insert_a0_chapter_before_delivery(document)
    _patch_delivery_section(document)

    document.save(str(OUTPUT_REPORT))

    DELIVERABLE.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUTPUT_REPORT, DELIVERABLE)

    print(f"源（未改动）: {SOURCE_REPORT}")
    print(f"已另存为: {OUTPUT_REPORT}")
    print(f"已同步: {DELIVERABLE}")


if __name__ == "__main__":
    main()
