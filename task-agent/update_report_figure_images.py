#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将技术报告中指定图题下的配图替换为大字号重绘版本。"""

from __future__ import annotations

import shutil
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

REPORT_PATH = Path(
    "/srv/supercare/比赛文档/1_技术报告_超级[GP-护士-照护员]智能体协同算法_含序贯贝叶斯_20260525.docx"
)
DELIVERABLE = Path("/srv/supercare/比赛文档/交付物") / REPORT_PATH.name
FIG_DIR = Path(__file__).resolve().parent / "output"

# 锚点段落关键字 -> 图片文件（替换锚点段落后第一个含 drawing 的段落）
# Word 嵌入宽度（英寸）：接近 A4 正文满宽，避免大图被缩得过小
EMBED_WIDTH_FULL = 7.35

FIGURE_MAP: List[Tuple[str, str, float]] = [
    ("如图3所示，长者健康计算图 DataAgent", "图3_DataAgent工作流.png", EMBED_WIDTH_FULL),
    ("如图5所示，三超循证知识基座", "图5_循证知识基座工作流.png", EMBED_WIDTH_FULL),
    ("图5-1 整体技术链路", "图5-1_整体技术链路_爱照护序贯贝叶斯超级GP.png", EMBED_WIDTH_FULL),
]


def _text(paragraph: Paragraph) -> str:
    return (paragraph.text or "").strip()


def _has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"))


def _remove_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _insert_picture_after(anchor: Paragraph, image_path: Path, width_inches: float) -> None:
    new_el = OxmlElement("w:p")
    anchor._element.addnext(new_el)
    new_p = Paragraph(new_el, anchor._parent)
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path.is_file():
        new_p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    else:
        new_p.add_run(f"（配图缺失：{image_path.name}）")


def replace_figures(doc_path: Path) -> int:
    doc = Document(str(doc_path))
    replaced = 0

    paragraphs = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        anchor_text = _text(paragraph)
        if not anchor_text:
            continue
        for keyword, filename, width in FIGURE_MAP:
            if keyword not in anchor_text:
                continue
            image_path = FIG_DIR / filename
            target: Optional[Paragraph] = None
            for offset in range(1, 4):
                if index + offset >= len(paragraphs):
                    break
                candidate = paragraphs[index + offset]
                if _has_drawing(candidate):
                    target = candidate
                    break
            if target is None:
                _insert_picture_after(paragraph, image_path, width)
            else:
                _remove_paragraph(target)
                _insert_picture_after(paragraph, image_path, width)
            replaced += 1
            break

    doc.save(str(doc_path))
    return replaced


def main() -> None:
    count = replace_figures(REPORT_PATH)
    DELIVERABLE.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REPORT_PATH, DELIVERABLE)
    print(f"已替换 {count} 处配图: {REPORT_PATH}")
    print(f"已同步: {DELIVERABLE}")


if __name__ == "__main__":
    main()
